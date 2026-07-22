"""Contract tests for supported RFP inputs and the canonical DOCX export.

Every document in this module is generated during the test.  No repository
fixture or production data directory is read or written.
"""
from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path


def _write_text_pdf(path: Path, page_texts: list[str]) -> None:
    """Generate a small, native-text PDF using the installed pypdf package."""
    from pypdf import PdfWriter
    from pypdf.generic import (
        DecodedStreamObject,
        DictionaryObject,
        NameObject,
    )

    writer = PdfWriter()
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
            NameObject("/Encoding"): NameObject("/WinAnsiEncoding"),
        }
    )
    font_ref = writer._add_object(font)

    for text in page_texts:
        page = writer.add_blank_page(width=612, height=792)
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): font_ref}
                )
            }
        )
        escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = DecodedStreamObject()
        stream.set_data(
            f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("cp1252")
        )
        page[NameObject("/Contents")] = writer._add_object(stream)

    writer.write(path)


def test_generated_pdf_extracts_native_text_with_page_boundaries(tmp_path: Path) -> None:
    from app.jobs.intake import _extract_text_for_intake
    from app.services.pdf_extract import extract_text_for_path

    path = tmp_path / "solicitation.PDF"
    _write_text_pdf(
        path,
        [
            "The contractor shall provide migration services.",
            "Proposals must include a transition schedule.",
        ],
    )

    text, page_count = extract_text_for_path(path)
    assert page_count == 2
    assert "--- Page 1 ---" in text
    assert "migration services" in text
    assert "--- Page 2 ---" in text
    assert "transition schedule" in text

    intake_text, intake_page_count = _extract_text_for_intake(str(path), path.name)
    assert intake_page_count == 2
    assert intake_text == text


def test_generated_docx_extracts_paragraphs_and_table_cells(tmp_path: Path) -> None:
    from docx import Document

    from app.jobs.intake import _extract_text_for_intake
    from app.services.pdf_extract import extract_text_for_path

    path = tmp_path / "statement-of-work.DOCX"
    document = Document()
    document.add_heading("Statement of Work", level=1)
    document.add_paragraph("The contractor shall operate the service desk.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Deliverable"
    table.cell(0, 1).text = "Due"
    table.cell(1, 0).text = "Transition plan"
    table.cell(1, 1).text = "Day 30"
    document.save(path)

    text, page_count = extract_text_for_path(path)
    assert page_count == 0
    assert "Statement of Work" in text
    assert "operate the service desk" in text
    assert "Deliverable | Due" in text
    assert "Transition plan | Day 30" in text

    intake_text, intake_page_count = _extract_text_for_intake(str(path), path.name)
    assert intake_page_count == 1
    assert intake_text.startswith("--- Page 1 ---\n")
    assert text in intake_text


def test_generated_xlsx_extracts_all_sheets_with_intake_boundaries(tmp_path: Path) -> None:
    from openpyxl import Workbook

    from app.jobs.intake import _extract_text_for_intake
    from app.services.pdf_extract import extract_text_for_path

    path = tmp_path / "requirements.XLSX"
    workbook = Workbook()
    scope = workbook.active
    scope.title = "Scope"
    scope.append(["Requirement", "Priority"])
    scope.append(["Provide 24x7 monitoring", "Mandatory"])
    pricing = workbook.create_sheet("Pricing")
    pricing.append(["Labor category", "Hours"])
    pricing.append(["Program Manager", 160])
    workbook.save(path)
    workbook.close()

    text, sheet_count = extract_text_for_path(path)
    assert sheet_count == 2
    assert "=== Sheet: Scope ===" in text
    assert "Provide 24x7 monitoring | Mandatory" in text
    assert "=== Sheet: Pricing ===" in text
    assert "Program Manager | 160" in text

    intake_text, intake_page_count = _extract_text_for_intake(str(path), path.name)
    assert intake_page_count == 2
    assert "--- Page 1 ---\n[Sheet: Scope]" in intake_text
    assert "--- Page 2 ---\n[Sheet: Pricing]" in intake_text
    assert "Program Manager | 160" in intake_text


def test_exported_docx_is_openable_and_preserves_draft_content(
    inmemory_db, monkeypatch, tmp_path: Path,
) -> None:
    import app.db.session as db_session
    import app.services.export as export_service
    import app.services.sections as section_service
    from app.core.enums import ProposalStatus
    from app.models import Proposal, ProposalSection, RfpPackage

    # ``sections`` imports session_scope directly, so patch its alias too in
    # case another test imported the module before the isolated DB fixture ran.
    monkeypatch.setattr(section_service, "session_scope", db_session.session_scope)
    monkeypatch.setattr(export_service, "extract_submission_filename", lambda _pid: None)

    with db_session.session_scope() as db:
        package = RfpPackage(
            uploaded_at=datetime.now(UTC),
            storage_dir="memory://format-contract",
        )
        db.add(package)
        db.flush()
        proposal = Proposal(
            rfp_package_id=package.id,
            title="Generated Export Contract",
            status=ProposalStatus.DRAFT_READY,
        )
        db.add(proposal)
        db.flush()
        proposal_id = proposal.id
        db.add(
            ProposalSection(
                proposal_id=proposal_id,
                section_id="SEC-001",
                section_title="Technical Approach",
                section_order=1,
                section_brief="Describe the tested technical approach.",
                draft_text_markdown=(
                    "### Delivery Model\n"
                    "We provide **verified continuity** through every transition.\n\n"
                    "- Named transition lead\n"
                    "- Weekly risk review\n\n"
                    "| Metric | Target |\n"
                    "| --- | --- |\n"
                    "| Service availability | 99.9% |"
                ),
                current_revision_number=3,
                compliance_items_addressed_json=[],
                citations_json=[],
                needs_human_placeholders_json=[],
                shortfall_mitigations_applied_json=[],
            )
        )

    data, filename, summary = export_service.compile_proposal_to_docx(
        proposal_id,
        proposal_title="Generated Export Contract",
        include_submission_checklist=False,
    )
    output_path = tmp_path / filename
    output_path.write_bytes(data)

    # Reopen the emitted artifact through python-docx.  This validates the
    # package relationships/XML, not merely that the result starts with ZIP.
    from docx import Document

    reopened = Document(output_path)
    paragraph_text = "\n".join(p.text for p in reopened.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in reopened.tables
        for row in table.rows
        for cell in row.cells
    )

    assert filename == "generated-export-contract.docx"
    assert len(data) == summary["byte_count"]
    assert summary["total_sections"] == 1
    assert summary["sections_skipped"] == []
    assert "Generated Export Contract" in paragraph_text
    assert "SEC-001" in paragraph_text
    assert "Technical Approach" in paragraph_text
    assert "verified continuity" in paragraph_text
    assert "Named transition lead" in paragraph_text
    assert "Metric" in table_text
    assert "Service availability" in table_text
    assert "99.9%" in table_text

    # Loading from an in-memory byte stream exercises the exact byte payload
    # returned to the UI download handler as well as the saved-file path above.
    assert Document(BytesIO(data)).paragraphs
