"""Isolated browser lifecycle for KB, profile, guidance, and decisions.

This module deliberately owns a second disposable application process instead
of using the session-scoped E2E workspace. Approving a profile suggestion is a
real filesystem mutation and must not leak into proposal-workflow tests that
expect the baseline synthetic profile.
"""
from __future__ import annotations

import json
import os
import re
import socket
import sqlite3
import subprocess
import sys
import tempfile
import time
import urllib.error
import urllib.request
import uuid
from collections.abc import Iterator
from contextlib import nullcontext
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import urlsplit

import pytest
from playwright.sync_api import expect

from tests.e2e.conftest import (
    PROJECT_ROOT,
    _is_benign_navigation_abort,
    _read_log,
    _seed_data_root,
    _unexpected_server_error_blocks,
    _write_browser_issues,
)

pytestmark = pytest.mark.e2e

_DOC_FILENAME = "synthetic_browser_capability.docx"
_CERTIFICATION = "E2E Browser Lifecycle Certification"
_INITIAL_RULE = "Avoid unsupported lifecycle claims in synthetic responses."
_EDITED_RULE = (
    "Tie every synthetic lifecycle claim to an explicitly named evidence artifact."
)
_DECISION_ID = "DEC-E2E-001"
_DECISION_TOPIC = "Synthetic browser lifecycle staffing"
_PROVIDER_KEYS = (
    "ANTHROPIC_API_KEY",
    "OPENAI_API_KEY",
    "GOOGLE_API_KEY",
    "GEMINI_API_KEY",
    "GROK_API_KEY",
    "VOYAGE_API_KEY",
)


@dataclass
class KbLifecycleSession:
    page: Any
    base_url: str
    root: Path
    database_path: Path
    artifacts: Path
    source_docx: Path
    initial_profile: dict[str, Any]
    rule_id: int


def _free_loopback_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def _write_source_docx(path: Path) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Synthetic Corporate Capability Statement", level=1)
    document.add_paragraph(
        "Synthetic E2E Company LLC provides firm-wide delivery governance, "
        "secure browser validation, and measurable lifecycle controls."
    )
    document.add_paragraph(
        f"The company holds the {_CERTIFICATION}. This generated document is "
        "test-only evidence stored exclusively in a disposable workspace."
    )
    document.save(path)


def _seed_rule_and_decision(database_path: Path, root: Path) -> int:
    with sqlite3.connect(database_path) as db:
        cursor = db.execute(
            """
            INSERT INTO learned_rules (
                kind, rule_text, source_action, source_category,
                source_severity, source_reviewer, status, hits
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                "writer_avoid",
                _INITIAL_RULE,
                "accept",
                "uncited_claim",
                "MINOR",
                None,
                "draft",
                0,
            ),
        )
        rule_id = int(cursor.lastrowid)
        db.commit()

    decisions_path = root / "decisions.json"
    decisions_path.write_text(
        json.dumps(
            {
                "_meta": {"version": "e2e-1.0.0"},
                "decisions": [
                    {
                        "id": _DECISION_ID,
                        "topic": _DECISION_TOPIC,
                        "decision": (
                            "Use a named synthetic validation owner for browser "
                            "acceptance."
                        ),
                        "applies_to_gaps_like": (
                            "Lifecycle staffing and acceptance ownership gaps."
                        ),
                        "established_on": "2030-01-02",
                        "source_proposal_id": None,
                        "source_gap_id": None,
                    }
                ],
            },
            indent=2,
            sort_keys=True,
        )
        + "\n",
        encoding="utf-8",
    )
    return rule_id


@pytest.fixture()
def isolated_kb_browser(
    request: pytest.FixtureRequest,
    e2e_enabled: None,
    playwright_browser: Any,
    tmp_path: Path,
) -> Iterator[KbLifecycleSession]:
    """Start a dedicated migrated app and isolated Chromium context."""
    root = (tmp_path / "kb-lifecycle-workspace").resolve()
    root.mkdir(parents=True)
    _seed_data_root(root)
    profile_path = root / "company_profile.json"
    initial_profile = json.loads(profile_path.read_text(encoding="utf-8"))

    artifact_base = Path(
        os.environ.get(
            "RFP_E2E_ARTIFACT_ROOT",
            str(Path(tempfile.gettempdir()) / "rfp-e2e-artifacts"),
        )
    ).resolve()
    artifacts = artifact_base / f"kb-lifecycle-{uuid.uuid4().hex[:10]}"
    artifacts.mkdir(parents=True, exist_ok=True)

    database_path = root / "sqlite.db"
    environment = {str(key): str(value) for key, value in os.environ.items()}
    environment.pop("PYTEST_CURRENT_TEST", None)
    environment.update(
        {
            "APP_ENV": "e2e",
            "APP_HOST": "127.0.0.1",
            "APP_STORAGE_SECRET": f"e2e-kb-{uuid.uuid4().hex}",
            "RFP_DATA_DIR": str(root),
            "DATABASE_URL": f"sqlite:///{database_path.as_posix()}",
            "RFP_E2E_FAKE_LLM": "1",
            "RFP_E2E_LLM_FIXTURES": str(
                PROJECT_ROOT / "tests" / "e2e" / "fixtures" / "workflow_llm.json"
            ),
            "RFP_E2E_LLM_LEDGER": str(artifacts / "llm_calls.jsonl"),
            "PYTHONUNBUFFERED": "1",
            "NO_PROXY": "127.0.0.1,localhost",
            "no_proxy": "127.0.0.1,localhost",
        }
    )
    for key in _PROVIDER_KEYS:
        environment[key] = ""

    migration = subprocess.run(
        [sys.executable, "-m", "alembic", "upgrade", "head"],
        cwd=PROJECT_ROOT,
        env=environment,
        capture_output=True,
        text=True,
        timeout=120,
        check=False,
    )
    migration_log = artifacts / "migration.log"
    migration_log.write_text(
        (migration.stdout or "") + (migration.stderr or ""), encoding="utf-8"
    )
    if migration.returncode != 0 or not database_path.is_file():
        pytest.fail(
            "isolated KB migration failed; "
            f"see {migration_log}",
            pytrace=False,
        )

    rule_id = _seed_rule_and_decision(database_path, root)
    source_docx = root / "upload_sources" / _DOC_FILENAME
    _write_source_docx(source_docx)

    port = _free_loopback_port()
    environment["APP_PORT"] = str(port)
    base_url = f"http://127.0.0.1:{port}"
    server_log = artifacts / "server.log"
    log_handle = server_log.open("w", encoding="utf-8")
    creation_flags = subprocess.CREATE_NEW_PROCESS_GROUP if os.name == "nt" else 0
    process = subprocess.Popen(
        [sys.executable, str(PROJECT_ROOT / "tests/e2e/support/run_app.py")],
        cwd=PROJECT_ROOT,
        env=environment,
        stdout=log_handle,
        stderr=subprocess.STDOUT,
        text=True,
        creationflags=creation_flags,
    )

    startup_error: str | None = None
    deadline = time.monotonic() + 30.0
    while time.monotonic() < deadline:
        if process.poll() is not None:
            startup_error = f"server exited early with code {process.returncode}"
            break
        try:
            with urllib.request.urlopen(
                f"{base_url}/api/health", timeout=1.0
            ) as response:
                payload = json.loads(response.read().decode("utf-8"))
            if response.status == 200 and payload.get("ok") is True:
                break
        except (OSError, ValueError, urllib.error.URLError):
            time.sleep(0.1)
    else:
        startup_error = "health endpoint did not become ready within 30 seconds"

    if startup_error is not None:
        if process.poll() is None:
            process.terminate()
            process.wait(timeout=8)
        log_handle.close()
        pytest.fail(
            f"isolated KB application startup failed: {startup_error}\n"
            f"{_read_log(server_log)}",
            pytrace=False,
        )

    issues: list[str] = []
    tearing_down = {"value": False}
    context = None
    page = None
    try:
        # Reuse the suite's Playwright runtime. Starting sync_playwright() here
        # fails when another E2E test has already activated that runtime.
        with nullcontext(playwright_browser) as browser:
            context = browser.new_context()

            def route_request(route: Any) -> None:
                parsed = urlsplit(route.request.url)
                if parsed.scheme in {"http", "https", "ws", "wss"} and (
                    parsed.hostname not in {"127.0.0.1", "localhost"}
                ):
                    issues.append(
                        f"blocked external browser request: {route.request.url}"
                    )
                    route.abort("blockedbyclient")
                    return
                route.continue_()

            context.route("**/*", route_request)
            context.tracing.start(screenshots=True, snapshots=True, sources=True)
            page = context.new_page()
            page.on(
                "console",
                lambda message: issues.append(
                    f"browser console error: {message.text}"
                )
                if message.type == "error"
                else None,
            )
            page.on(
                "pageerror",
                lambda error: issues.append(f"uncaught page error: {error}"),
            )

            def on_request_failed(failed_request: Any) -> None:
                if tearing_down["value"]:
                    return
                if _is_benign_navigation_abort(failed_request):
                    return
                if urlsplit(failed_request.url).hostname in {
                    "127.0.0.1",
                    "localhost",
                }:
                    issues.append(
                        f"local request failed: {failed_request.method} "
                        f"{failed_request.url} ({failed_request.failure})"
                    )

            page.on("requestfailed", on_request_failed)
            page.on(
                "response",
                lambda response: issues.append(
                    f"HTTP {response.status}: {response.request.method} "
                    f"{response.url}"
                )
                if response.status >= 500
                else None,
            )

            yield KbLifecycleSession(
                page=page,
                base_url=base_url,
                root=root,
                database_path=database_path,
                artifacts=artifacts,
                source_docx=source_docx,
                initial_profile=initial_profile,
                rule_id=rule_id,
            )

            failed = bool(
                getattr(request.node, "rep_call", None)
                and request.node.rep_call.failed
            )
            if failed or issues:
                try:
                    page.screenshot(
                        path=str(artifacts / "failure.png"), full_page=True
                    )
                    (artifacts / "page.html").write_text(
                        page.content(), encoding="utf-8"
                    )
                except Exception as exc:
                    issues.append(f"failed to capture failure artifacts: {exc}")
                _write_browser_issues(artifacts, issues, test_failed=failed)
            tearing_down["value"] = True
            try:
                page.goto("about:blank", wait_until="commit", timeout=3_000)
                page.wait_for_timeout(100)
            except Exception:
                pass
            context.tracing.stop(path=str(artifacts / "trace.zip"))
            context.close()
            context = None
            if issues and not failed:
                pytest.fail(
                    "Unexpected isolated KB browser errors:\n- "
                    + "\n- ".join(issues)
                    + f"\nArtifacts: {artifacts}",
                    pytrace=False,
                )
    finally:
        if context is not None:
            try:
                context.close()
            except Exception:
                pass
        if process.poll() is None:
            process.terminate()
            try:
                process.wait(timeout=8)
            except subprocess.TimeoutExpired:
                process.kill()
                process.wait(timeout=5)
        log_handle.close()
        backend_errors = _unexpected_server_error_blocks(_read_log(server_log))
        if backend_errors:
            pytest.fail(
                "Unexpected isolated KB backend errors:\n\n"
                + "\n\n".join(backend_errors)
                + f"\n\nServer log: {server_log}",
                pytrace=False,
            )


def _row(database_path: Path, sql: str, params: tuple = ()):
    with sqlite3.connect(database_path) as db:
        return db.execute(sql, params).fetchone()


def _wait_for_row(
    database_path: Path,
    sql: str,
    params: tuple,
    predicate,
    *,
    timeout: float = 20.0,
):
    deadline = time.monotonic() + timeout
    last = None
    while time.monotonic() < deadline:
        last = _row(database_path, sql, params)
        if predicate(last):
            return last
        time.sleep(0.05)
    raise AssertionError(f"database condition was not met; last row={last!r}")


def _wait_for_json(path: Path, predicate, *, timeout: float = 10.0) -> dict:
    deadline = time.monotonic() + timeout
    last: dict[str, Any] | None = None
    while time.monotonic() < deadline:
        try:
            last = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            time.sleep(0.05)
            continue
        if predicate(last):
            return last
        time.sleep(0.05)
    raise AssertionError(f"JSON condition was not met; last payload={last!r}")


def _goto(page: Any, url: str, expected_text: str) -> None:
    response = page.goto(url, wait_until="domcontentloaded")
    assert response is not None and response.status == 200
    page.get_by_text(expected_text, exact=False).filter(visible=True).first.wait_for(
        state="visible", timeout=15_000
    )


def _visible_card(page: Any, text: str):
    return page.locator(".q-card").filter(has_text=text).filter(visible=True).first


def _open_guidance_rule(
    page: Any,
    base_url: str,
    text: str,
    *,
    archived: bool = False,
):
    _goto(page, f"{base_url}/kb", "Knowledge Base")
    page.get_by_role("tab", name="Learned Guidance", exact=True).click()
    if archived:
        page.get_by_text(re.compile(r"^Archived \(\d+\)$")).filter(
            visible=True
        ).click()
    page.get_by_text(text, exact=True).filter(visible=True).wait_for(
        state="visible", timeout=10_000
    )
    return _visible_card(page, text)


def test_kb_profile_guidance_and_decision_lifecycle(
    isolated_kb_browser: KbLifecycleSession,
) -> None:
    session = isolated_kb_browser
    page = session.page
    base = session.base_url
    database_path = session.database_path
    profile_path = session.root / "company_profile.json"

    # Real DOCX upload -> deterministic LLM classification -> async extraction
    # and fact extraction. The database and managed file prove each persisted
    # boundary completed, rather than relying only on transient notifications.
    _goto(page, f"{base}/kb", "Knowledge Base")
    upload_input = page.locator('input[type="file"]')
    assert upload_input.count() == 1
    upload_input.set_input_files(str(session.source_docx))
    staged_row = page.get_by_role(
        "listitem", name=f"Staged KB file {_DOC_FILENAME}", exact=True
    )
    staged_row.wait_for(state="visible", timeout=10_000)
    classification = staged_row.get_by_text(
        "corporate (confidence: high)", exact=False
    ).filter(
        visible=True
    )
    classification.wait_for(state="visible", timeout=15_000)
    expect(staged_row.get_by_text(_DOC_FILENAME, exact=True)).to_be_visible()
    page.get_by_role("button", name="Save to KB", exact=True).click()

    document_row = _wait_for_row(
        database_path,
        """
        SELECT id, storage_path, document_class, status, tags_json,
               extracted_text_md
        FROM knowledge_base_documents WHERE filename = ?
        """,
        (_DOC_FILENAME,),
        lambda row: bool(row and row[3] == "active" and row[5]),
    )
    document_id = int(document_row[0])
    storage_path = Path(document_row[1]).resolve()
    assert document_row[2] == "corporate"
    assert set(json.loads(document_row[4])) == {
        "synthetic-kb-lifecycle",
        "browser-e2e",
    }
    assert "Synthetic Corporate Capability Statement" in document_row[5]
    managed_root = (session.root / "kb_documents").resolve()
    assert storage_path.is_relative_to(managed_root)
    assert storage_path.parent == managed_root / str(document_id)
    assert storage_path.read_bytes() == session.source_docx.read_bytes()

    # Saving introduces the document-list expansion component dynamically.
    # Exercise that refreshed UI before navigating away: otherwise a fast fake
    # ingestion can let the next page.goto cancel expansion.js in flight and
    # turn an intentional navigation race into a false transport failure.
    page.get_by_role("button", name="Expand all", exact=True).click()
    saved_document = page.locator(f'[data-kb-document-id="{document_id}"]')
    expect(saved_document).to_be_visible(timeout=10_000)
    expect(saved_document.get_by_text(_DOC_FILENAME, exact=True)).to_be_visible()

    suggestion_row = _wait_for_row(
        database_path,
        """
        SELECT id, status, summary, proposed_value_json
        FROM profile_suggestions
        WHERE kb_document_id = ? AND section = 'certifications'
        """,
        (document_id,),
        lambda row: bool(row and row[1] == "pending"),
    )
    suggestion_id = int(suggestion_row[0])
    assert suggestion_row[2] == f"Add certification: {_CERTIFICATION}"
    assert json.loads(suggestion_row[3]) == _CERTIFICATION

    ledger_entries = [
        json.loads(line)
        for line in (session.artifacts / "llm_calls.jsonl")
        .read_text(encoding="utf-8")
        .splitlines()
        if line.strip()
    ]
    assert [entry["agent_name"] for entry in ledger_entries] == [
        "kb_classify",
        "kb_facts_corporate",
    ]

    # Approving is the real profile write path: status changes, the file is
    # atomically rewritten with a new version, and the server cache reloads.
    _goto(page, f"{base}/config?tab=suggestions", "Pending Profile Updates")
    summary = f"Add certification: {_CERTIFICATION}"
    page.get_by_text(summary, exact=True).filter(visible=True).wait_for()
    _visible_card(page, summary).get_by_role(
        "button", name="Approve", exact=True
    ).click()
    page.get_by_text("No pending profile updates.", exact=True).wait_for(
        state="visible", timeout=10_000
    )
    _wait_for_row(
        database_path,
        "SELECT status FROM profile_suggestions WHERE id = ?",
        (suggestion_id,),
        lambda row: row == ("approved",),
    )
    initial_version = session.initial_profile["_meta"]["version"]
    updated_profile = _wait_for_json(
        profile_path,
        lambda payload: (
            payload.get("_meta", {}).get("version") != initial_version
            and _CERTIFICATION in payload.get("certifications", [])
        ),
    )
    updated_version = updated_profile["_meta"]["version"]

    _goto(page, f"{base}/config?tab=profile", f"Version {updated_version}")
    page.get_by_text("Certifications", exact=True).filter(visible=True).click()
    page.get_by_text(_CERTIFICATION, exact=False).filter(visible=True).wait_for()

    # Learned Guidance: edit -> approve -> archive -> re-approve -> archive ->
    # delete. Each step is re-opened from a fresh page to avoid depending on
    # expansion state while the tab's five-second refresh timer is active.
    rule_card = _open_guidance_rule(page, base, _INITIAL_RULE)
    rule_card.get_by_role("button", name="Edit", exact=True).click()
    edit_dialog = page.get_by_role("dialog")
    edit_dialog.get_by_text("Edit rule", exact=True).wait_for()
    edit_dialog.get_by_label("Rule text", exact=True).fill(_EDITED_RULE)
    edit_dialog.get_by_role("button", name="Save", exact=True).click()
    _wait_for_row(
        database_path,
        "SELECT rule_text, status FROM learned_rules WHERE id = ?",
        (session.rule_id,),
        lambda row: row == (_EDITED_RULE, "draft"),
    )

    rule_card = _open_guidance_rule(page, base, _EDITED_RULE)
    rule_card.get_by_role("button", name="Approve", exact=True).click()
    _wait_for_row(
        database_path,
        "SELECT status FROM learned_rules WHERE id = ?",
        (session.rule_id,),
        lambda row: row == ("approved",),
    )

    rule_card = _open_guidance_rule(page, base, _EDITED_RULE)
    rule_card.get_by_role("button", name="Archive", exact=True).click()
    _wait_for_row(
        database_path,
        "SELECT status FROM learned_rules WHERE id = ?",
        (session.rule_id,),
        lambda row: row == ("archived",),
    )

    rule_card = _open_guidance_rule(
        page, base, _EDITED_RULE, archived=True
    )
    rule_card.get_by_role("button", name="Re-approve", exact=True).click()
    _wait_for_row(
        database_path,
        "SELECT status FROM learned_rules WHERE id = ?",
        (session.rule_id,),
        lambda row: row == ("approved",),
    )

    rule_card = _open_guidance_rule(page, base, _EDITED_RULE)
    rule_card.get_by_role("button", name="Archive", exact=True).click()
    _wait_for_row(
        database_path,
        "SELECT status FROM learned_rules WHERE id = ?",
        (session.rule_id,),
        lambda row: row == ("archived",),
    )
    rule_card = _open_guidance_rule(
        page, base, _EDITED_RULE, archived=True
    )
    rule_card.get_by_role("button", name="Delete", exact=True).click()
    _wait_for_row(
        database_path,
        "SELECT id FROM learned_rules WHERE id = ?",
        (session.rule_id,),
        lambda row: row is None,
    )

    # Decisions uses the real JSON-backed confirmation/delete workflow.
    _goto(page, f"{base}/config?tab=decisions", _DECISION_TOPIC)
    decision_card = _visible_card(page, _DECISION_TOPIC)
    decision_card.get_by_role(
        "button",
        name=f"Delete decision {_DECISION_ID}: {_DECISION_TOPIC}",
        exact=True,
    ).click()
    delete_decision_dialog = page.get_by_role("dialog")
    delete_decision_dialog.get_by_text(
        f"Delete decision {_DECISION_ID}?", exact=True
    ).wait_for()
    delete_decision_dialog.get_by_role(
        "button", name="Delete", exact=True
    ).click()
    page.get_by_text("No decisions recorded yet.", exact=False).wait_for(
        state="visible", timeout=10_000
    )
    decisions_payload = _wait_for_json(
        session.root / "decisions.json",
        lambda payload: payload.get("decisions") == [],
    )
    assert decisions_payload["_meta"]["version"] == "e2e-1.0.0"

    # Delete the KB source through its confirmation dialog. The row, managed
    # file/directory, and suggestion audit row are removed, while the approved
    # profile fact and bumped version deliberately remain.
    _goto(page, f"{base}/kb", "Knowledge Base")
    page.get_by_role("button", name="Expand all", exact=True).filter(
        visible=True
    ).click()
    document_ui_row = page.locator(
        f'[data-kb-document-id="{document_id}"]'
    )
    expect(document_ui_row).to_be_visible(timeout=10_000)
    expect(document_ui_row.get_by_text(_DOC_FILENAME, exact=True)).to_be_visible()
    document_ui_row.get_by_role(
        "button", name=f"Delete KB document #{document_id}", exact=True
    ).click()
    delete_document_dialog = page.get_by_role("dialog")
    delete_document_dialog.get_by_text(
        f"Delete KB document #{document_id}?", exact=True
    ).wait_for()
    delete_document_dialog.get_by_role(
        "button", name="Delete", exact=True
    ).click()

    _wait_for_row(
        database_path,
        "SELECT id FROM knowledge_base_documents WHERE id = ?",
        (document_id,),
        lambda row: row is None,
    )
    assert _row(
        database_path,
        "SELECT id FROM profile_suggestions WHERE id = ?",
        (suggestion_id,),
    ) is None
    assert not storage_path.exists()
    assert not storage_path.parent.exists()

    surviving_profile = json.loads(profile_path.read_text(encoding="utf-8"))
    assert surviving_profile["_meta"]["version"] == updated_version
    assert _CERTIFICATION in surviving_profile["certifications"]
