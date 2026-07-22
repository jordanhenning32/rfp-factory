"""Real-browser proof of the first proposal workflow and its human gates."""
from __future__ import annotations

import json
import re
import sqlite3
import time
from collections.abc import Callable
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from tests.e2e.conftest import BrowserSession

pytestmark = pytest.mark.e2e

_TITLE = "Synthetic Modernization Services"
_AGENCY = "E2E Department of Technology"
_FILENAME = "synthetic_modernization_rfp.docx"
_REQ_1 = (
    "The contractor shall design and implement a secure cloud-hosted case "
    "management platform."
)
_REQ_2 = (
    "The offeror must provide a project management and quality assurance approach."
)
_EXPECTED_LLM_CALLS = {
    ("complete", "intake_metadata", None),
    ("call_tool", "compliance_matrix", "report_compliance_items"),
    ("call_tool", "compliance_completeness", "report_completeness_review"),
    ("call_tool", "compliance_validator", "report_validation_results"),
    ("call_tool", "section_m_extractor", "report_evaluation_criteria"),
    ("call_tool", "shortfall_strategist", "report_gap_analyses"),
    ("call_tool", "outline_agent", "report_proposal_outline"),
}


def _make_synthetic_rfp(path: Path) -> None:
    """Build an original DOCX fixture inside this run's disposable workspace."""
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Synthetic Modernization Services RFP", level=0)
    document.add_paragraph(
        "SYNTHETIC E2E TEST DOCUMENT — this is not a real solicitation."
    )
    document.add_paragraph(f"Issuing Agency: {_AGENCY}")
    document.add_paragraph("Solicitation Number: E2E-RFP-001")
    document.add_paragraph("NAICS: 541511")
    document.add_paragraph("Proposal Due Date: December 15, 2030")
    document.add_heading("3.1 Technical Approach", level=1)
    document.add_paragraph(_REQ_1)
    document.add_heading("3.2 Management Approach", level=1)
    document.add_paragraph(_REQ_2)
    document.add_heading("4. Evaluation Criteria", level=1)
    document.add_paragraph("Factor 1 — Technical Approach carries a weight of 60 percent.")
    document.add_paragraph(
        "Responses to Requirements REQ-001 and REQ-002 will be evaluated under "
        "Factor 1."
    )
    document.add_paragraph("Award will be made on a best-value basis.")
    document.save(path)


def _fetchone(database_path: Path, sql: str, params: tuple[Any, ...] = ()) -> sqlite3.Row | None:
    with sqlite3.connect(database_path, timeout=2.0) as db:
        db.row_factory = sqlite3.Row
        return db.execute(sql, params).fetchone()


def _fetchall(database_path: Path, sql: str, params: tuple[Any, ...] = ()) -> list[sqlite3.Row]:
    with sqlite3.connect(database_path, timeout=2.0) as db:
        db.row_factory = sqlite3.Row
        return db.execute(sql, params).fetchall()


def _eventually(
    probe: Callable[[], Any],
    predicate: Callable[[Any], bool],
    *,
    description: str,
    timeout: float = 45.0,
) -> Any:
    deadline = time.monotonic() + timeout
    last_value: Any = None
    last_error: Exception | None = None
    while time.monotonic() < deadline:
        try:
            last_value = probe()
            last_error = None
            if predicate(last_value):
                return last_value
        except sqlite3.OperationalError as exc:
            last_error = exc
        time.sleep(0.1)
    detail = f"last value={last_value!r}"
    if last_error is not None:
        detail += f", last SQLite error={last_error!r}"
    raise AssertionError(f"Timed out waiting for {description}; {detail}")


def _wait_for_status(database_path: Path, proposal_id: int, expected: str) -> sqlite3.Row:
    row = _eventually(
        lambda: _fetchone(
            database_path,
            "SELECT id, status FROM proposals WHERE id = ?",
            (proposal_id,),
        ),
        lambda value: value is not None and value["status"] == expected,
        description=f"proposal #{proposal_id} status={expected!r}",
    )
    assert isinstance(row, sqlite3.Row)
    return row


def _ledger_calls(path: Path) -> list[tuple[str, str, str | None]]:
    entries = [
        json.loads(line)
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    return [
        (entry["method"], entry["agent_name"], entry.get("tool_name"))
        for entry in entries
    ]


def test_new_proposal_through_outline_approval(browser_session: BrowserSession) -> None:
    from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
    from playwright.sync_api import expect

    page = browser_session.page
    server = browser_session.server
    workspace = server.workspace
    database_path = workspace.database_path

    source_path = workspace.root / "source-inputs" / _FILENAME
    _make_synthetic_rfp(source_path)

    response = page.goto(
        f"{server.base_url}/proposals/new",
        wait_until="domcontentloaded",
    )
    assert response is not None and response.status == 200
    page.get_by_text("Upload an RFP package", exact=True).wait_for(
        state="visible", timeout=10_000,
    )

    upload_input = page.locator('input[type="file"]')
    assert upload_input.count() == 1
    upload_input.set_input_files(str(source_path))

    page.get_by_role(
        "listitem", name=f"Staged RFP file {_FILENAME}", exact=True
    ).wait_for(state="visible", timeout=10_000)
    title_input = page.get_by_label("Proposal title", exact=True)
    expect(title_input).to_have_value(_TITLE, timeout=15_000)
    expect(page.get_by_label("Agency", exact=True)).to_have_value(_AGENCY)
    expect(page.get_by_label("NAICS", exact=True)).to_have_value("541511")
    expect(page.get_by_label("Due date", exact=True)).to_have_value("2030-12-15")
    expect(page.get_by_label("Notes (optional)", exact=True)).to_have_value(
        "Solicitation #: E2E-RFP-001"
    )
    page.get_by_text(re.compile(r"^✓ Auto-filled:"), exact=False).wait_for(
        state="visible", timeout=10_000,
    )

    page.get_by_role("button", name="Run", exact=True).click()
    page.wait_for_url(re.compile(r"/proposals/\d+/progress$"), timeout=15_000)
    match = re.search(r"/proposals/(\d+)/progress$", page.url)
    assert match is not None
    proposal_id = int(match.group(1))

    _wait_for_status(database_path, proposal_id, "awaiting_scope_signoff")

    expect(page.get_by_text("Status: Awaiting scope sign-off", exact=True)).to_be_visible(
        timeout=10_000
    )
    expect(page.get_by_text(_TITLE, exact=True)).to_be_visible()
    expect(page.get_by_text(f"Agency: {_AGENCY}", exact=True)).to_be_visible()
    expect(page.get_by_text("NAICS: 541511", exact=True)).to_be_visible()
    expect(page.get_by_text("Due: 2030-12-15", exact=True)).to_be_visible()
    expect(page.get_by_text(re.compile(r"^RFP package .* 1 file\(s\)$"))).to_be_visible()
    expect(page.get_by_text(_FILENAME, exact=True)).to_be_visible()
    expect(page.get_by_text("2 package items", exact=True)).to_be_visible()
    expect(
        page.get_by_text(
            f"{_FILENAME} — Complete",
            exact=True,
        )
    ).to_be_visible()
    expect(page.get_by_text("reviewer Google · gemini-2.5-pro", exact=False)).to_be_visible()
    expect(page.get_by_text("Compliance matrix: 2 item(s) extracted.", exact=True)).to_be_visible()
    expect(page.get_by_text("Evaluation criteria: 1 factor(s) extracted.", exact=True)).to_be_visible()
    expect(page.get_by_text("Shortfall analysis: 0 gap(s) flagged.", exact=True)).to_be_visible()
    expect(page.get_by_role("button", name="Open Proposal Review", exact=True)).to_be_visible()

    proposal = _fetchone(
        database_path,
        """
        SELECT p.title, p.agency, p.naics, p.due_date, p.role, p.status,
               p.notes, p.service_line, p.evaluation_criteria_json,
               r.storage_dir
        FROM proposals AS p
        JOIN rfp_packages AS r ON r.id = p.rfp_package_id
        WHERE p.id = ?
        """,
        (proposal_id,),
    )
    assert proposal is not None
    assert proposal["title"] == _TITLE
    assert proposal["agency"] == _AGENCY
    assert proposal["naics"] == "541511"
    assert proposal["due_date"] == "2030-12-15"
    assert proposal["role"] == "prime"
    assert proposal["status"] == "awaiting_scope_signoff"
    assert proposal["notes"] == "Solicitation #: E2E-RFP-001"
    assert proposal["service_line"] == "it_services"

    evaluation = json.loads(proposal["evaluation_criteria_json"])
    assert evaluation["evaluation_method"] == "best_value"
    assert evaluation["trade_off_language"] == "Award will be made on a best-value basis."
    assert evaluation["section_l_to_m_map"] == {
        "REQ-001": ["F1"],
        "REQ-002": ["F1"],
    }
    assert evaluation["factors"][0]["factor_name"] == "Technical Approach"
    assert evaluation["factors"][0]["weight_pct"] == 60

    document = _fetchone(
        database_path,
        """
        SELECT d.filename, d.storage_path, d.page_count, d.extracted_text_md,
               d.structure_json
        FROM rfp_package_documents AS d
        JOIN proposals AS p ON p.rfp_package_id = d.rfp_package_id
        WHERE p.id = ?
        """,
        (proposal_id,),
    )
    assert document is not None
    assert document["filename"] == _FILENAME
    # The intake dispatcher normalizes page-less DOCX extraction to one
    # logical page so package summaries never report a zero-page document.
    assert document["page_count"] == 1
    assert _REQ_1 in document["extracted_text_md"]
    assert _REQ_2 in document["extracted_text_md"]
    assert json.loads(document["structure_json"])["content_sha256"]

    workspace_root = workspace.root.resolve()
    package_dir = Path(proposal["storage_dir"]).resolve()
    stored_path = Path(document["storage_path"]).resolve()
    assert package_dir.is_relative_to(workspace_root)
    assert stored_path.is_relative_to(package_dir)
    assert package_dir.is_dir()
    assert stored_path.is_file()
    assert stored_path.read_bytes() == source_path.read_bytes()

    compliance = _fetchall(
        database_path,
        """
        SELECT requirement_id, requirement_text, source_doc, source_section,
               source_page, requirement_type, category, status
        FROM compliance_matrix_items
        WHERE proposal_id = ?
        ORDER BY requirement_id
        """,
        (proposal_id,),
    )
    assert [row["requirement_id"] for row in compliance] == ["REQ-001", "REQ-002"]
    assert [row["requirement_text"] for row in compliance] == [_REQ_1, _REQ_2]
    assert [row["requirement_type"] for row in compliance] == ["shall", "must"]
    assert [row["category"] for row in compliance] == ["technical", "management"]
    assert all(row["source_doc"] == _FILENAME for row in compliance)
    assert all(row["status"] == "active" for row in compliance)
    assert _fetchone(
        database_path,
        "SELECT COUNT(*) AS count FROM gap_analyses WHERE proposal_id = ?",
        (proposal_id,),
    )["count"] == 0

    stage_rows = _fetchall(
        database_path,
        """
        SELECT agent_name, status, error_text
        FROM agent_runs
        WHERE proposal_id = ? AND agent_name = '_stage'
        ORDER BY id
        """,
        (proposal_id,),
    )
    assert stage_rows
    assert all(row["agent_name"] == "_stage" for row in stage_rows)
    assert all(row["status"] == "completed" for row in stage_rows)
    stage_messages = [row["error_text"] for row in stage_rows]
    assert any(message == "Parsed 1 document(s)." for message in stage_messages)
    assert any(message == "Compliance matrix: 2 item(s) extracted." for message in stage_messages)
    assert any(message == "Evaluation criteria: 1 factor(s) extracted." for message in stage_messages)
    assert any(message == "Shortfall analysis: 0 gap(s) flagged." for message in stage_messages)

    review_url = f"{server.base_url}/proposals/{proposal_id}"
    review_deadline = time.monotonic() + 15.0
    while page.url != review_url and time.monotonic() < review_deadline:
        # The progress page refreshes itself while its final pipeline events
        # arrive.  If that refresh replaces the button during a click, retry
        # the same user action against the newly rendered control.
        page.get_by_role("button", name="Open Proposal Review", exact=True).click()
        try:
            page.wait_for_url(review_url, timeout=3_000, wait_until="domcontentloaded")
        except PlaywrightTimeoutError:
            pass
    assert page.url == review_url
    expect(
        page.get_by_text("Action needed: review gaps and sign off scope", exact=True)
    ).to_be_visible(timeout=10_000)
    page.get_by_role("button", name="Sign off scope", exact=True).click()
    _wait_for_status(database_path, proposal_id, "drafting")
    expect(
        page.get_by_text("Scope signed off — generate the section outline", exact=True)
    ).to_be_visible(timeout=10_000)

    page.get_by_role("button", name="Generate Draft Outline", exact=True).click()
    page.wait_for_url(f"{server.base_url}/proposals/{proposal_id}/progress", timeout=10_000)
    _wait_for_status(database_path, proposal_id, "awaiting_outline_approval")
    expect(page.get_by_text("Status: Awaiting outline approval", exact=True)).to_be_visible(
        timeout=10_000
    )
    expect(page.get_by_text("Outline ready: 2 section(s).", exact=True)).to_be_visible()

    outline_rows = _fetchall(
        database_path,
        """
        SELECT section_id, section_title, section_order, page_limit, word_limit,
               requires_cost_analysis, compliance_items_addressed_json,
               draft_text_markdown, current_revision_number
        FROM proposal_sections
        WHERE proposal_id = ?
        ORDER BY section_order
        """,
        (proposal_id,),
    )
    assert [row["section_id"] for row in outline_rows] == ["SEC-001", "SEC-002"]
    assert [row["section_title"] for row in outline_rows] == [
        "Technical Approach",
        "Project Management and Quality Assurance",
    ]
    assert [json.loads(row["compliance_items_addressed_json"]) for row in outline_rows] == [
        ["REQ-001"],
        ["REQ-002"],
    ]
    assert [row["word_limit"] for row in outline_rows] == [2000, 1000]
    assert all(row["requires_cost_analysis"] == 0 for row in outline_rows)
    assert all(row["draft_text_markdown"] is None for row in outline_rows)
    assert all(row["current_revision_number"] == 0 for row in outline_rows)

    page.goto(f"{server.base_url}/proposals/{proposal_id}", wait_until="domcontentloaded")
    expect(
        page.get_by_text("Action needed: review the outline and approve", exact=True)
    ).to_be_visible(timeout=10_000)
    expect(page.get_by_text("2 sections", exact=True)).to_be_visible()
    expect(page.get_by_text("2 of 2 narrative items mapped", exact=True)).to_be_visible()
    expect(page.get_by_text("SEC-001", exact=True)).to_be_visible()
    expect(page.get_by_text("SEC-002", exact=True)).to_be_visible()
    page.get_by_role("button", name="Approve Outline", exact=True).click()

    _wait_for_status(database_path, proposal_id, "awaiting_team_approval")
    expect(
        page.get_by_text("Action needed: build and approve the team", exact=True)
    ).to_be_visible(timeout=10_000)
    expect(page.get_by_text("Team Composition", exact=True)).to_be_visible()
    expect(page.get_by_role("button", name="Propose Team (AI)", exact=True)).to_be_visible()

    for tab_name in (
        "Compliance",
        "Evaluation Criteria",
        "Gaps",
        "Outline",
        "Team",
        "Cost",
        "Draft",
        "Reviewer Findings",
        "Submission Checklist",
        "Timeline",
        "Spend",
    ):
        assert page.get_by_role("tab", name=tab_name, exact=True).count() == 1

    page.get_by_role("tab", name="Outline", exact=True).click()
    expect(page.get_by_text("Outline approved", exact=True)).to_be_visible()
    expect(page.get_by_text("2 sections", exact=True)).to_be_visible()

    ledger_path = workspace.artifacts / "llm_calls.jsonl"
    assert ledger_path.is_file()
    calls = [
        call for call in _ledger_calls(ledger_path)
        if call in _EXPECTED_LLM_CALLS
    ]
    assert len(calls) == len(_EXPECTED_LLM_CALLS)
    assert set(calls) == _EXPECTED_LLM_CALLS
