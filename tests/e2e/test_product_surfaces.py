"""Real-browser coverage for the complete persisted product surface.

The test intentionally uses a representative graph rather than mocks at the
UI boundary: the real NiceGUI application, real SQLite services, real export
code, and a real Chromium browser are involved. Only LLM providers are the
deterministic fixtures installed by the E2E server process.
"""
from __future__ import annotations

import json
import re
import sqlite3
import subprocess
import sys
import time
from pathlib import Path
from zipfile import ZipFile

import pytest
from playwright.sync_api import expect

from tests.e2e.conftest import BrowserSession, E2EServer

pytestmark = pytest.mark.e2e

PROJECT_ROOT = Path(__file__).resolve().parents[2]
PROPOSAL_TABS = (
    "Compliance",
    "Evaluation Criteria",
    "Win Strategy",
    "Amendments & Q&A",
    "Gaps",
    "Outline",
    "Team",
    "Cost",
    "Cost Review",
    "Draft",
    "Reviewer Findings",
    "Final Polish",
    "Completed Draft",
    "Submission Checklist",
    "Timeline",
    "Spend",
)


@pytest.fixture()
def surface_seed(e2e_server: E2EServer):
    """Populate only the migrated disposable database used by E2E."""
    result = subprocess.run(
        [
            sys.executable,
            str(
                PROJECT_ROOT
                / "tests"
                / "e2e"
                / "support"
                / "seed_surface_data.py"
            ),
        ],
        cwd=PROJECT_ROOT,
        env=e2e_server.workspace.environment,
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )
    seed_log = e2e_server.workspace.artifacts / "surface_seed.log"
    seed_log.write_text(
        (result.stdout or "") + (result.stderr or ""), encoding="utf-8"
    )
    if result.returncode != 0:
        pytest.fail(
            f"synthetic surface seed failed (exit {result.returncode}); "
            f"see {seed_log}",
            pytrace=False,
        )
    try:
        payload = json.loads((result.stdout or "").strip().splitlines()[-1])
    except (IndexError, ValueError) as exc:
        pytest.fail(
            f"synthetic surface seed returned invalid JSON: {exc}; see {seed_log}",
            pytrace=False,
        )
    assert payload["it_proposal_id"] > 0
    assert payload["payment_proposal_id"] > 0
    yield payload

    cleanup = subprocess.run(
        [
            sys.executable,
            str(
                PROJECT_ROOT
                / "tests"
                / "e2e"
                / "support"
                / "seed_surface_data.py"
            ),
            "--cleanup",
        ],
        cwd=PROJECT_ROOT,
        env=e2e_server.workspace.environment,
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )
    cleanup_log = e2e_server.workspace.artifacts / "surface_cleanup.log"
    cleanup_log.write_text(
        (cleanup.stdout or "") + (cleanup.stderr or ""), encoding="utf-8"
    )
    if cleanup.returncode != 0:
        pytest.fail(
            f"synthetic surface cleanup failed (exit {cleanup.returncode}); "
            f"see {cleanup_log}",
            pytrace=False,
        )


@pytest.fixture()
def seeded_browser(
    surface_seed: dict[str, int],
    browser_session: BrowserSession,
) -> tuple[BrowserSession, dict[str, int]]:
    # This dependency wrapper guarantees seeding finishes before Chromium opens
    # its first product route without changing the shared conftest harness.
    return browser_session, surface_seed


def _goto(page, url: str, expected_text: str) -> None:
    response = page.goto(url, wait_until="domcontentloaded")
    assert response is not None
    assert response.status == 200
    page.get_by_text(expected_text, exact=False).filter(visible=True).first.wait_for(
        state="visible", timeout=15_000
    )


def _proposal_tab(page, label: str):
    # Material icons and action badges are descendants of the q-tab and both
    # contribute to ``innerText`` in Chromium. Match the dedicated label node
    # instead of the entire tab's text/accessibility name.
    label_node = page.locator(".q-tab__label").filter(
        has_text=re.compile(rf"^\s*{re.escape(label)}\s*$", re.I)
    )
    return page.get_by_role("tab").filter(has=label_node).first


def _open_proposal_tab(page, label: str, expected_text: str) -> None:
    tab = _proposal_tab(page, label)
    tab.wait_for(state="visible", timeout=10_000)
    tab.click()
    page.get_by_text(expected_text, exact=False).filter(visible=True).first.wait_for(
        state="visible", timeout=10_000
    )
    expect(tab).to_have_attribute("aria-selected", "true", timeout=10_000)


def _row(db_path: Path, sql: str, params: tuple = ()):
    with sqlite3.connect(db_path) as db:
        return db.execute(sql, params).fetchone()


def _wait_for_row(
    db_path: Path,
    sql: str,
    params: tuple,
    predicate,
    *,
    timeout: float = 6.0,
):
    deadline = time.monotonic() + timeout
    last = None
    while time.monotonic() < deadline:
        last = _row(db_path, sql, params)
        if predicate(last):
            return last
        time.sleep(0.05)
    raise AssertionError(f"database condition was not met; last row={last!r}")


def _card_with_text(page, text: str):
    return page.locator(".q-card").filter(has_text=text).first


def test_all_routes_tabs_persisted_actions_export_and_archive_read_only(
    seeded_browser: tuple[BrowserSession, dict[str, int]],
) -> None:
    session, seed = seeded_browser
    page = session.page
    server = session.server
    base = server.base_url
    db_path = server.workspace.database_path
    it_id = seed["it_proposal_id"]
    payment_id = seed["payment_proposal_id"]

    # Route 1: Pipeline. Both service-line branches and the outcome summary are
    # represented in the actual proposal cards.
    _goto(page, base, "Proposals in flight")
    page.get_by_text("Synthetic IT Modernization RFP", exact=True).wait_for()
    page.get_by_text("Synthetic Payment Processing RFP", exact=True).wait_for()
    page.get_by_text("Win rate:", exact=False).wait_for()

    # Route 2: New Proposal. Upload mechanics receive their own focused test;
    # this suite verifies the complete real route and service-line control load.
    _goto(page, f"{base}/proposals/new", "Upload an RFP package")
    page.get_by_text("Drop RFP files here", exact=True).wait_for()
    page.get_by_text("Service line", exact=True).wait_for()

    # Route 3: Run Progress with a populated package and run ledger.
    _goto(
        page,
        f"{base}/proposals/{it_id}/progress",
        "Synthetic IT Modernization RFP",
    )
    page.get_by_text("synthetic_it_rfp.txt", exact=True).wait_for()
    page.get_by_text("Pipeline activity", exact=True).wait_for()
    page.get_by_text("4 package items", exact=True).wait_for(timeout=10_000)

    # Route 4: Proposal Review, all sixteen tabs. The checks use distinct body
    # content so they prove the selected panel rendered, not merely its tab.
    _goto(
        page,
        f"{base}/proposals/{it_id}",
        "Synthetic IT Modernization RFP",
    )
    tab_expectations = (
        ("Compliance", "The offeror shall provide a phased modernization approach."),
        ("Evaluation Criteria", "Evaluation Factors"),
        ("Win Strategy", "Evaluator scorecard, win themes, proof"),
        ("Amendments & Q&A", "Upload Amendment"),
        ("Gaps", "GAP-001"),
        ("Outline", "Exclude from draft"),
        ("Team", "Jordan Example"),
        ("Cost", "Cost pipeline"),
        (
            "Cost Review",
            "Security validation hours may be understated for the synthetic scope.",
        ),
        ("Draft", "The unique browser-test sentence"),
        ("Reviewer Findings", "Clarify how transition-gate evidence is verified."),
        ("Final Polish", "Cross-section consistency cleanup"),
        ("Completed Draft", "The unique browser-test sentence"),
        ("Submission Checklist", "Provide the synthetic transition-gate evidence matrix."),
        ("Timeline", "Seeded Discovery"),
        ("Spend", "By pipeline stage"),
    )
    for tab_label, expected in tab_expectations:
        _open_proposal_tab(page, tab_label, expected)

    # Win Strategy: run the deterministic persisted refresh through Chromium.
    _open_proposal_tab(page, "Win Strategy", "Win Strategy")
    with page.expect_navigation(wait_until="domcontentloaded", timeout=15_000):
        page.get_by_role("button", name="Generate All", exact=True).click()
    page.get_by_text("Synthetic IT Modernization RFP", exact=True).wait_for(
        state="visible", timeout=15_000
    )
    strategy_row = _wait_for_row(
        db_path,
        """
        SELECT evaluator_scorecard_json, win_themes_json,
               past_performance_matches_json, price_to_win_json,
               red_team_findings_json, graphics_tables_json
        FROM proposals WHERE id = ?
        """,
        (it_id,),
        lambda row: bool(row and all(row)),
    )
    assert all(json.loads(value) for value in strategy_row)

    # Gaps: choose a mitigation and prove the service persisted it.
    _open_proposal_tab(page, "Gaps", "GAP-001")
    page.get_by_role("button", name="Choose this", exact=True).first.click()
    page.get_by_text("YOUR CHOICE", exact=True).wait_for(timeout=10_000)
    _wait_for_row(
        db_path,
        "SELECT selected_mitigation_index FROM gap_analyses WHERE id = ?",
        (seed["gap_id"],),
        lambda row: row == (0,),
    )

    # Outline: exercise the persisted skip switch in both directions so later
    # draft/export coverage still includes every section.
    _open_proposal_tab(
        page,
        "Outline",
        "Exclude from draft",
    )
    outline_switch = page.get_by_role("switch", name="Exclude from draft").first
    if outline_switch.count() == 0:
        outline_switch = page.get_by_role(
            "checkbox", name="Exclude from draft"
        ).first
    outline_switch.click()
    _wait_for_row(
        db_path,
        "SELECT COUNT(*) FROM proposal_sections WHERE proposal_id = ? AND excluded_from_draft = 1",
        (it_id,),
        lambda row: row == (1,),
    )
    outline_switch = page.get_by_role("switch", name="Exclude from draft").first
    if outline_switch.count() == 0:
        outline_switch = page.get_by_role(
            "checkbox", name="Exclude from draft"
        ).first
    outline_switch.click()
    _wait_for_row(
        db_path,
        "SELECT COUNT(*) FROM proposal_sections WHERE proposal_id = ? AND excluded_from_draft = 1",
        (it_id,),
        lambda row: row == (0,),
    )

    # Team: re-approval is a safe persisted roster action and exercises the
    # complete approval gate without invoking the AI team composer.
    _open_proposal_tab(page, "Team", "Team Composition")
    prior_approval = _row(
        db_path, "SELECT team_approved_at FROM proposals WHERE id = ?", (it_id,)
    )[0]
    page.get_by_role("button", name="Re-approve Team", exact=True).click()
    _wait_for_row(
        db_path,
        "SELECT team_approved_at FROM proposals WHERE id = ?",
        (it_id,),
        lambda row: bool(row and row[0] and row[0] != prior_approval),
    )

    # Cost: click the real scenario card and verify the selected proposal
    # posture changed from MEDIUM to HIGH.
    _open_proposal_tab(page, "Cost", "Proposed price")
    page.get_by_text("High — Protective", exact=True).filter(
        visible=True
    ).first.click()
    page.get_by_text("High — Protective — detail", exact=True).filter(
        visible=True
    ).wait_for(timeout=10_000)
    _wait_for_row(
        db_path,
        "SELECT proposed_scenario FROM proposals WHERE id = ?",
        (it_id,),
        lambda row: row == ("HIGH",),
    )

    # Cost Review: accept the relational finding (no agent call).
    _open_proposal_tab(
        page,
        "Cost Review",
        "Security validation hours may be understated for the synthetic scope.",
    )
    cost_card = _card_with_text(
        page, "Security validation hours may be understated for the synthetic scope."
    )
    cost_card.get_by_role("button", name="Accept", exact=True).click()
    cost_card.get_by_text("ACCEPTED", exact=True).wait_for(timeout=10_000)
    _wait_for_row(
        db_path,
        "SELECT user_action FROM cost_review_findings WHERE id = ?",
        (seed["cost_finding_id"],),
        lambda row: row == ("accepted",),
    )

    # Draft: make a manual source edit through the real textarea/save path.
    _open_proposal_tab(page, "Draft", "The unique browser-test sentence")
    technical_card = _card_with_text(page, "The unique browser-test sentence")
    technical_card.get_by_role("button", name="Edit", exact=True).click()
    technical_card = _card_with_text(page, "Editing markdown source")
    draft_editor = technical_card.get_by_role("textbox").first
    edited_sentence = "Browser-saved revision marker 2030."
    draft_editor.fill(draft_editor.input_value() + f"\n\n{edited_sentence}")
    technical_card.get_by_role("button", name="Save changes", exact=True).click()
    page.get_by_text(edited_sentence, exact=True).wait_for(timeout=10_000)
    _wait_for_row(
        db_path,
        "SELECT draft_text_markdown, current_revision_number FROM proposal_sections WHERE proposal_id = ? AND section_id = 'SEC-001'",
        (it_id,),
        lambda row: bool(row and edited_sentence in row[0] and row[1] == 3),
    )

    # Reviewer Findings: unmark an accepted item. This deliberately avoids the
    # accept path's asynchronous lesson-extraction LLM call.
    _open_proposal_tab(
        page,
        "Reviewer Findings",
        "Clarify how transition-gate evidence is verified.",
    )
    finding_card = _card_with_text(
        page, "Clarify how transition-gate evidence is verified."
    )
    finding_card.get_by_role("button", name="Unmark", exact=True).click()
    finding_card.get_by_role("button", name="Accept", exact=True).wait_for(
        timeout=10_000
    )
    _wait_for_row(
        db_path,
        "SELECT accepted_at, dismissed_at FROM reviewer_findings WHERE id = ?",
        (seed["reviewer_finding_id"],),
        lambda row: row == (None, None),
    )

    # Submission Checklist: toggle the matrix-backed form item.
    _open_proposal_tab(
        page,
        "Submission Checklist",
        "Submit the signed synthetic representations form.",
    )
    checklist_card = _card_with_text(
        page, "Submit the signed synthetic representations form."
    )
    checklist_card.get_by_role("checkbox").click()
    _wait_for_row(
        db_path,
        "SELECT submission_obtained FROM compliance_matrix_items WHERE id = ?",
        (seed["it_requirement_id"],),
        lambda row: row == (1,),
    )

    # Timeline: add a real phase via dialog and verify the serialized document.
    _open_proposal_tab(page, "Timeline", "Implementation Timeline")
    page.get_by_role("button", name="Add Phase", exact=True).click()
    dialog = page.get_by_role("dialog")
    dialog.get_by_text("Add Phase", exact=True).wait_for()
    dialog.get_by_label("Phase name", exact=True).fill("Browser Validation")
    dialog.get_by_label("Deliverable / output", exact=True).fill(
        "Validated E2E acceptance report"
    )
    dialog.get_by_label("Owner / role (optional)", exact=True).fill("QA Lead")
    dialog.get_by_role("button", name="Add", exact=True).click()
    page.get_by_text("Browser Validation", exact=True).first.wait_for(timeout=10_000)
    timeline_row = _wait_for_row(
        db_path,
        "SELECT timeline_json FROM proposals WHERE id = ?",
        (it_id,),
        lambda row: bool(row and "Browser Validation" in (row[0] or "")),
    )
    assert len(json.loads(timeline_row[0])["phases"]) == 2

    # Completed Draft: download the actual export and validate it as both an
    # OPC ZIP package and a Word document with the freshly persisted edit.
    _open_proposal_tab(page, "Completed Draft", "Completed Draft")
    # The Completed Draft panel is mounted with the rest of the proposal page;
    # use its explicit refresh control after an in-page manual draft edit.
    page.get_by_role("button", name="Refresh", exact=True).filter(
        visible=True
    ).click()
    page.get_by_text(edited_sentence, exact=True).filter(visible=True).wait_for(
        timeout=10_000
    )
    with page.expect_download(timeout=15_000) as download_info:
        page.get_by_role("button", name="Download DOCX", exact=True).click()
    download = download_info.value
    docx_path = session.artifacts / download.suggested_filename
    download.save_as(str(docx_path))
    assert docx_path.suffix.lower() == ".docx"
    assert docx_path.stat().st_size > 5_000
    with ZipFile(docx_path) as archive:
        names = set(archive.namelist())
        assert "[Content_Types].xml" in names
        assert "word/document.xml" in names
        assert archive.testzip() is None
    from docx import Document

    document = Document(str(docx_path))
    exported_text = "\n".join(p.text for p in document.paragraphs)
    exported_text += "\n" + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Synthetic IT Modernization RFP" in exported_text
    assert edited_sentence in exported_text
    assert "Submit the signed synthetic representations form" in exported_text

    # Payment Systems branch: real scan, model override, JSON-backed Cost
    # Review action, and outcome upsert all run through the browser.
    _goto(
        page,
        f"{base}/proposals/{payment_id}",
        "Synthetic Payment Processing RFP",
    )
    page.get_by_text("Outcome", exact=True).wait_for()
    page.get_by_role("button", name="Edit outcome", exact=True).click()
    outcome_dialog = page.get_by_role("dialog")
    outcome_dialog.get_by_text("Edit Outcome", exact=True).wait_for()
    awarded_to = outcome_dialog.get_by_role("textbox", name="Awarded to")
    if awarded_to.count() == 0:
        awarded_to = outcome_dialog.get_by_role("textbox").nth(2)
    awarded_to.fill("Browser-Updated Synthetic Awardee")
    outcome_dialog.get_by_role("button", name="Save", exact=True).click()
    page.get_by_text("Browser-Updated Synthetic Awardee", exact=False).wait_for(
        timeout=15_000
    )
    _wait_for_row(
        db_path,
        "SELECT awarded_to FROM proposal_outcomes WHERE proposal_id = ?",
        (payment_id,),
        lambda row: row == ("Browser-Updated Synthetic Awardee",),
    )

    _open_proposal_tab(page, "Cost", "Payment Market Scan")
    page.get_by_text("Recommended Bid Posture", exact=True).wait_for()
    page.get_by_text("SyntheticPay", exact=True).first.wait_for()
    page.get_by_role("button", name="Flat rate", exact=True).filter(
        visible=True
    ).click()
    _wait_for_row(
        db_path,
        "SELECT selected_pricing_model FROM proposals WHERE id = ?",
        (payment_id,),
        lambda row: row == ("flat_rate",),
    )
    page.get_by_text("Synthetic Payment Processing RFP", exact=True).wait_for(
        timeout=15_000
    )
    _open_proposal_tab(page, "Cost Review", "Cost Review (Payment Systems)")
    payment_finding_card = _card_with_text(
        page, "The chargeback fee needs an explicit when-incurred qualifier."
    )
    payment_finding_card.get_by_role("button", name="Accept", exact=True).click()
    payment_finding_card.get_by_text("ACCEPTED", exact=True).wait_for(
        timeout=10_000
    )
    payment_review_row = _wait_for_row(
        db_path,
        "SELECT payment_cost_review_findings_json FROM proposals WHERE id = ?",
        (payment_id,),
        lambda row: bool(row and '"user_action": "accepted"' in row[0]),
    )
    assert json.loads(payment_review_row[0])["findings"][0]["user_action"] == "accepted"

    # Route 5: Knowledge Base, including a persisted learned-rule approval.
    _goto(page, f"{base}/kb", "Knowledge Base")
    page.get_by_role("button", name="Expand all", exact=True).filter(
        visible=True
    ).click()
    page.get_by_text("synthetic_capability.txt", exact=True).filter(
        visible=True
    ).wait_for()
    kb_guidance_tab = page.get_by_role("tab", name="Learned Guidance", exact=True)
    kb_guidance_tab.click()
    page.get_by_text(
        "Tie each transition claim to a named acceptance artifact.", exact=True
    ).wait_for()
    rule_card = _card_with_text(
        page, "Tie each transition claim to a named acceptance artifact."
    )
    rule_card.get_by_role("button", name="Approve", exact=True).click()
    _wait_for_row(
        db_path,
        "SELECT status FROM learned_rules WHERE source_finding_id = ?",
        (seed["reviewer_finding_id"],),
        lambda row: row == ("approved",),
    )
    page.get_by_role("tab", name="Documents", exact=True).click()
    page.get_by_text("synthetic_capability.txt", exact=True).filter(
        visible=True
    ).wait_for()

    # Route 6: every Configuration panel, plus rejection of a pending profile
    # suggestion. Reject is used so the test never rewrites the profile fixture.
    _goto(page, f"{base}/config?tab=suggestions", "Pending Profile Updates")
    page.get_by_text("Add the synthetic delivery certification.", exact=True).wait_for()
    suggestion_card = _card_with_text(
        page, "Add the synthetic delivery certification."
    )
    suggestion_card.get_by_role("button", name="Reject", exact=True).click()
    page.get_by_text("No pending profile updates.", exact=True).wait_for(
        timeout=10_000
    )
    _wait_for_row(
        db_path,
        "SELECT status FROM profile_suggestions WHERE id = ?",
        (seed["profile_suggestion_id"],),
        lambda row: row == ("rejected",),
    )
    config_tabs = (
        ("Company Profile", "Synthetic E2E Company LLC"),
        ("Decisions", "Synthetic transition specialist"),
        ("Pricing Rules", "Pricing rules are active"),
        ("Models", "Per-agent model assignments are active"),
        ("Cost Caps", "Per-run and monthly cost-cap values"),
    )
    for tab_name, text in config_tabs:
        page.get_by_role("tab", name=tab_name, exact=True).click()
        page.get_by_text(text, exact=False).first.wait_for(
            state="visible", timeout=10_000
        )

    # Route 7: Admin summary reflects the representative seed graph.
    _goto(page, f"{base}/admin", "This is a read-only installation summary")
    page.get_by_text("Proposal status reference", exact=True).wait_for()
    page.get_by_text("e2e-1.0.0", exact=True).wait_for()

    # Archived contract: take the submitted Payment Systems proposal through
    # the real confirmation workflow, then prove all 16 tab bodies remain
    # readable while representative mutation controls are disabled.
    _goto(
        page,
        f"{base}/proposals/{payment_id}",
        "Synthetic Payment Processing RFP",
    )
    page.get_by_role("button", name="Archive proposal", exact=True).click()
    archive_dialog = page.get_by_role("dialog")
    archive_dialog.get_by_text("Archive proposal?", exact=True).wait_for()
    with page.expect_navigation(wait_until="domcontentloaded", timeout=15_000):
        archive_dialog.get_by_role("button", name="Archive", exact=True).click()
    page.get_by_text("Archived (read-only)", exact=True).wait_for(timeout=15_000)
    _wait_for_row(
        db_path,
        "SELECT status FROM proposals WHERE id = ?",
        (payment_id,),
        lambda row: row == ("archived",),
    )
    assert page.get_by_role(
        "button", name="Edit outcome", exact=True
    ).is_disabled()

    for label in PROPOSAL_TABS:
        tab = _proposal_tab(page, label)
        tab.click()
        expect(tab).to_have_attribute("aria-selected", "true", timeout=10_000)
        page.locator(".q-tab-panel").filter(visible=True).first.wait_for(
            state="visible", timeout=10_000
        )

    _open_proposal_tab(page, "Timeline", "Implementation Timeline")
    assert page.get_by_role("button", name="Add Phase", exact=True).is_disabled()

    _open_proposal_tab(page, "Cost", "Payment Market Scan")
    assert page.get_by_role("button", name="Flat rate", exact=True).filter(
        visible=True
    ).is_disabled()

    _open_proposal_tab(
        page,
        "Cost Review",
        "The chargeback fee needs an explicit when-incurred qualifier.",
    )
    assert page.locator(".q-tab-panel").filter(visible=True).get_by_role(
        "button", name="Reset to pending", exact=True
    ).is_disabled()

    # Reviewer Findings is refreshable and owns a five-second polling timer.
    # A delayed assertion guards against a re-render replacing the one-time
    # archived controls with newly active mutation buttons.
    _open_proposal_tab(
        page,
        "Reviewer Findings",
        "Confirm the payment-fee qualifier against the final schedule.",
    )
    archived_accept = _card_with_text(
        page, "Confirm the payment-fee qualifier against the final schedule."
    ).get_by_role("button", name="Accept", exact=True)
    assert archived_accept.is_disabled()
    page.wait_for_timeout(6_200)
    assert archived_accept.is_disabled()
