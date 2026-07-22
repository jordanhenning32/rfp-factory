"""Real-browser proof of immediate discovery and later matrix attachment."""
from __future__ import annotations

import re
import shutil
import sqlite3
import subprocess
import sys
import time
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from playwright.sync_api import expect

from tests.e2e.conftest import BrowserSession

pytestmark = pytest.mark.e2e


@pytest.fixture()
def isolated_cost_matrix_session(
    browser_session: BrowserSession,
) -> BrowserSession:
    """Restore every session-scoped resource this workflow mutates.

    The browser server and workspace intentionally live for the whole E2E
    session.  This test creates a real proposal and appends fake-provider
    ledger entries, so leaving either behind makes later empty-state and seed
    tests order-dependent.
    """
    workspace = browser_session.server.workspace
    with sqlite3.connect(workspace.database_path, timeout=2.0) as db:
        proposal_ids_before = {
            int(row[0]) for row in db.execute("SELECT id FROM proposals")
        }
        package_ids_before = {
            int(row[0]) for row in db.execute("SELECT id FROM rfp_packages")
        }

    ledger_path = workspace.artifacts / "llm_calls.jsonl"
    ledger_existed = ledger_path.exists()
    ledger_before = ledger_path.read_bytes() if ledger_existed else b""
    browser_input_dir = workspace.root / "cost-matrix-browser-inputs"

    try:
        yield browser_session
    finally:
        # Release the proposal page before removing the records it displays.
        try:
            browser_session.page.goto("about:blank", wait_until="commit")
        except Exception:
            pass

        with sqlite3.connect(workspace.database_path, timeout=2.0) as db:
            created_proposal_ids = [
                int(row[0])
                for row in db.execute("SELECT id FROM proposals")
                if int(row[0]) not in proposal_ids_before
            ]

        cleanup = subprocess.run(
            [
                sys.executable,
                "-c",
                (
                    "import sys\n"
                    "from app.db.session import session_scope\n"
                    "from app.services.proposals import delete_proposal\n"
                    "with session_scope() as db:\n"
                    "    for raw in sys.argv[1:]:\n"
                    "        result = delete_proposal(db, int(raw))\n"
                    "        if not result.get('deleted'):\n"
                    "            raise RuntimeError(f'cleanup failed for proposal {raw}: {result}')\n"
                ),
                *[str(proposal_id) for proposal_id in created_proposal_ids],
            ],
            cwd=Path(__file__).resolve().parents[2],
            env=workspace.environment,
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
        )

        # These are browser-side upload fixtures, not managed package files.
        # Resolve and verify the exact test directory before removing it.
        expected_input_dir = (
            workspace.root.resolve() / "cost-matrix-browser-inputs"
        )
        if (
            browser_input_dir.exists()
            and browser_input_dir.resolve() == expected_input_dir
        ):
            shutil.rmtree(browser_input_dir)

        if ledger_existed:
            ledger_path.write_bytes(ledger_before)
        else:
            ledger_path.unlink(missing_ok=True)

        assert cleanup.returncode == 0, cleanup.stdout + cleanup.stderr
        with sqlite3.connect(workspace.database_path, timeout=2.0) as db:
            assert {
                int(row[0]) for row in db.execute("SELECT id FROM proposals")
            } == proposal_ids_before
            assert {
                int(row[0]) for row in db.execute("SELECT id FROM rfp_packages")
            } == package_ids_before


def _make_rfp(path: Path) -> None:
    doc = Document()
    doc.add_heading("Synthetic Cost Matrix Solicitation", level=0)
    doc.add_paragraph("Issuing Agency: E2E Pricing Department")
    doc.add_paragraph("Solicitation Number: E2E-CM-001")
    doc.add_paragraph("NAICS: 541511")
    doc.add_paragraph("Proposal Due Date: December 15, 2030")
    # Match the deterministic shared extraction fixture used by the isolated
    # E2E server. Source-grounded production extraction intentionally rejects
    # fixture requirements that do not occur in the uploaded solicitation.
    doc.add_heading("3.1 Technical Approach", level=1)
    doc.add_paragraph(
        "The contractor shall design and implement a secure cloud-hosted "
        "case management platform."
    )
    doc.add_heading("3.2 Management Approach", level=1)
    doc.add_paragraph(
        "The offeror must provide a project management and quality assurance "
        "approach."
    )
    doc.add_paragraph(
        "The contractor shall perform Workstream Alpha and Workstream Beta."
    )
    doc.add_paragraph(
        "The offeror must provide a technical and project management approach."
    )
    doc.add_paragraph("Award will be made on a best-value basis.")
    doc.save(path)


def _make_matrix(path: Path, *, fee_schedule: bool = False) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Fee Schedule" if fee_schedule else "Buyer Pricing "
    ws["B2"] = "Vendor Name:"
    ws.merge_cells("C2:E2")
    ws["B3"] = "Date:"
    ws.merge_cells("C3:E3")
    ws["C3"].number_format = "m/d/yyyy"
    if fee_schedule:
        ws["B7"] = "Service"
        ws["C7"] = "Unit Fee"
        ws["B8"] = "Card transaction"
        ws["C8"] = 0
        ws["C8"].number_format = "$0.0000"
        ws["B9"] = "ACH transaction"
        ws["C9"] = 0
        ws["C9"].number_format = "$0.0000"
    else:
        ws["B8"] = "Work Item"
        ws["C8"] = "Total Cost"
        ws["B9"] = "Workstream Alpha"
        ws["C9"] = 0
        ws["C9"].number_format = "$#,##0.00"
        ws["B10"] = "Workstream Beta"
        ws["C10"] = 0
        ws["C10"].number_format = "$#,##0.00"
        ws["B11"] = "Total"
        ws["C11"] = "=SUM(C9:C10)"
        ws["C11"].number_format = "$#,##0.00"
    hidden = wb.create_sheet("Buyer source data")
    hidden.sheet_state = "hidden"
    hidden["A1"] = "Preserve this hidden source"
    wb.save(path)
    wb.close()


def _eventually(probe, predicate, *, timeout: float = 45.0):
    deadline = time.monotonic() + timeout
    last = None
    while time.monotonic() < deadline:
        try:
            last = probe()
            if predicate(last):
                return last
        except sqlite3.OperationalError:
            pass
        time.sleep(0.1)
    raise AssertionError(f"database condition not met; last={last!r}")


def _row(database: Path, sql: str, params: tuple = ()):
    with sqlite3.connect(database, timeout=2.0) as db:
        db.row_factory = sqlite3.Row
        return db.execute(sql, params).fetchone()


def _scalar(database: Path, sql: str, params: tuple = ()):
    row = _row(database, sql, params)
    return row[0] if row else None


def _cost_tab(page):
    label = page.locator(".q-tab__label").filter(
        has_text=re.compile(r"^\s*Cost\s*$", re.I)
    )
    return page.get_by_role("tab").filter(has=label).first


def test_cost_matrix_is_detected_before_intake_and_can_be_added_later(
    isolated_cost_matrix_session: BrowserSession,
) -> None:
    page = isolated_cost_matrix_session.page
    server = isolated_cost_matrix_session.server
    workspace = server.workspace
    source_dir = workspace.root / "cost-matrix-browser-inputs"
    source_dir.mkdir(parents=True, exist_ok=True)
    matrix_path = source_dir / "Buyer Cost Matrix.xlsx"
    rfp_path = source_dir / "synthetic_cost_matrix_rfp.docx"
    late_path = source_dir / "Late Fee Schedule.xlsx"
    _make_matrix(matrix_path)
    _make_rfp(rfp_path)
    _make_matrix(late_path, fee_schedule=True)

    response = page.goto(
        f"{server.base_url}/proposals/new",
        wait_until="domcontentloaded",
    )
    assert response is not None and response.status == 200
    upload = page.locator('input[type="file"]')
    assert upload.count() == 1

    # Upload the matrix first. It must be recognized without consuming the
    # solicitation metadata-extraction opportunity.
    upload.set_input_files(str(matrix_path))
    page.get_by_text("Cost matrix — held for pricing", exact=True).wait_for(
        state="visible",
        timeout=15_000,
    )
    expect(page.get_by_label("Proposal title", exact=True)).to_have_value("")

    upload.set_input_files(str(rfp_path))
    expect(page.get_by_label("Proposal title", exact=True)).to_have_value(
        "Synthetic Modernization Services",
        timeout=15_000,
    )
    expect(page.get_by_label("Agency", exact=True)).to_have_value(
        "E2E Department of Technology"
    )

    page.get_by_role("button", name="Run", exact=True).click()
    page.wait_for_url(re.compile(r"/proposals/\d+/progress$"), timeout=15_000)
    proposal_match = re.search(r"/proposals/(\d+)/progress$", page.url)
    assert proposal_match is not None
    proposal_id = int(proposal_match.group(1))

    _eventually(
        lambda: _scalar(
            workspace.database_path,
            "SELECT COUNT(*) FROM cost_matrix_artifacts WHERE proposal_id = ?",
            (proposal_id,),
        ),
        lambda value: value == 1,
    )
    _eventually(
        lambda: _scalar(
            workspace.database_path,
            "SELECT status FROM proposals WHERE id = ?",
            (proposal_id,),
        ),
        lambda value: value == "awaiting_scope_signoff",
    )
    matrix_doc = _row(
        workspace.database_path,
        """
        SELECT d.document_role, d.document_type, d.extracted_text_md
        FROM cost_matrix_artifacts AS a
        JOIN rfp_package_documents AS d ON d.id = a.source_document_id
        WHERE a.proposal_id = ?
        """,
        (proposal_id,),
    )
    assert matrix_doc is not None
    assert matrix_doc["document_role"] == "cost_matrix"
    assert matrix_doc["document_type"] == "form_template"
    assert matrix_doc["extracted_text_md"] is None

    page.goto(
        f"{server.base_url}/proposals/{proposal_id}",
        wait_until="domcontentloaded",
    )
    cost_tab = _cost_tab(page)
    cost_tab.click()
    panel = page.locator('[data-testid="cost-matrix-panel"]')
    expect(panel).to_be_visible(timeout=15_000)
    expect(panel.get_by_text("Buyer Cost Matrix.xlsx", exact=True)).to_be_visible()
    expect(panel.get_by_text("Buyer total rule detected", exact=False)).to_be_visible()
    expect(panel.get_by_role("button", name="Generate completed copy", exact=True)).to_be_disabled()

    late_upload = panel.locator('input[type="file"]')
    assert late_upload.count() == 1
    late_upload.set_input_files(str(late_path))
    expect(panel.get_by_text("Late Fee Schedule.xlsx", exact=True)).to_be_visible(
        timeout=15_000
    )
    _eventually(
        lambda: _scalar(
            workspace.database_path,
            "SELECT COUNT(*) FROM cost_matrix_artifacts WHERE proposal_id = ?",
            (proposal_id,),
        ),
        lambda value: value == 2,
    )
    late_analysis = _scalar(
        workspace.database_path,
        """
        SELECT a.analysis_json
        FROM cost_matrix_artifacts AS a
        JOIN rfp_package_documents AS d ON d.id = a.source_document_id
        WHERE a.proposal_id = ? AND d.filename = ?
        """,
        (proposal_id, "Late Fee Schedule.xlsx"),
    )
    assert late_analysis is not None and "Card transaction" in late_analysis
