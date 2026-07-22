"""Seed an isolated payment-systems proposal at its cost-workflow gate.

Upstream intake, outline, and team approval are deliberate fixture seams. The
browser test drives every payment-specific cost step and the shared downstream
review/submission workflow through production UI controls and jobs.
"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
from datetime import UTC, date, datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

TITLE = "Synthetic Payment Workflow RFP"
PACKAGE_SLUG = "synthetic-payment-workflow"


def _validate_environment() -> Path:
    if os.environ.get("APP_ENV", "").strip().lower() != "e2e":
        raise RuntimeError("payment seed requires APP_ENV=e2e")
    if os.environ.get("RFP_E2E_FAKE_LLM", "") != "1":
        raise RuntimeError("payment seed requires RFP_E2E_FAKE_LLM=1")

    raw_data_dir = os.environ.get("RFP_DATA_DIR", "").strip()
    if not raw_data_dir:
        raise RuntimeError("payment seed requires an explicit RFP_DATA_DIR")
    data_dir = Path(raw_data_dir).resolve()
    canonical_data = (PROJECT_ROOT / "data").resolve()
    try:
        data_dir.relative_to(canonical_data)
    except ValueError:
        pass
    else:
        raise RuntimeError(f"refusing to seed canonical data: {data_dir}")

    expected_db = (data_dir / "sqlite.db").as_posix()
    database_url = os.environ.get("DATABASE_URL", "").replace("\\", "/")
    if database_url != f"sqlite:///{expected_db}":
        raise RuntimeError(
            "payment seed DATABASE_URL must point exactly to sqlite.db "
            f"inside RFP_DATA_DIR; got {database_url!r}"
        )
    if not (data_dir / "sqlite.db").is_file():
        raise RuntimeError("payment seed requires a migrated E2E database")

    profile_path = data_dir / "company_profile.json"
    try:
        profile = json.loads(profile_path.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:
        raise RuntimeError("payment seed requires the synthetic profile") from exc
    if not str((profile.get("_meta") or {}).get("version") or "").startswith(
        "e2e-"
    ):
        raise RuntimeError("payment seed rejected a non-E2E profile")
    return data_dir


def _cleanup(data_dir: Path) -> dict[str, int | bool]:
    from sqlalchemy import select

    from app.db.session import SessionLocal
    from app.models import AgentRun, Proposal, RfpPackage

    deleted = 0
    with SessionLocal() as db:
        proposals = db.execute(
            select(Proposal).where(Proposal.title == TITLE)
        ).scalars().all()
        package_ids = [proposal.rfp_package_id for proposal in proposals]
        for proposal in proposals:
            db.query(AgentRun).filter(
                AgentRun.proposal_id == proposal.id
            ).delete(synchronize_session=False)
            db.delete(proposal)
            deleted += 1
        db.flush()
        for package_id in package_ids:
            package = db.get(RfpPackage, package_id)
            if package is not None:
                db.delete(package)
        db.commit()

    package_dir = data_dir / "rfp_packages" / PACKAGE_SLUG
    if package_dir.exists():
        shutil.rmtree(package_dir)
    return {"deleted": deleted, "ok": True}


def _seed(data_dir: Path) -> dict[str, int]:
    from docx import Document
    from sqlalchemy import select

    from app.core.enums import (
        ComplianceStatus,
        ProposalRole,
        ProposalStatus,
        RequirementCategory,
        RequirementType,
        RfpDocumentType,
    )
    from app.db.session import SessionLocal
    from app.models import (
        ComplianceMatrixItem,
        Proposal,
        ProposalSection,
        ProposalTeamMember,
        RfpPackage,
        RfpPackageDocument,
    )

    package_dir = data_dir / "rfp_packages" / PACKAGE_SLUG
    source_path = package_dir / "synthetic_payment_rfp.docx"
    package_dir.mkdir(parents=True, exist_ok=True)
    source_text = (
        "SYNTHETIC E2E SOLICITATION. The processor shall provide secure card "
        "and ACH payment acceptance, recurring billing, U.S.-only data "
        "residency, and an auditable fee schedule. The offeror shall describe "
        "implementation governance and provide all proposed transaction fees."
    )
    document = Document()
    document.add_heading("Synthetic Payment Workflow RFP", level=0)
    document.add_paragraph("SYNTHETIC E2E SOLICITATION — NOT A LIVE RFP")
    document.add_heading("Payment Processing Requirements", level=1)
    document.add_paragraph(source_text)
    document.save(source_path)

    now = datetime.now(UTC)
    with SessionLocal() as db:
        existing = db.execute(
            select(Proposal).where(Proposal.title == TITLE)
        ).scalar_one_or_none()
        if existing is not None:
            return {"proposal_id": existing.id}

        package = RfpPackage(
            uploaded_by="e2e-payment-seed",
            uploaded_at=now,
            storage_dir=str(package_dir),
            notes="Synthetic package for payment workflow coverage.",
        )
        db.add(package)
        db.flush()
        db.add(
            RfpPackageDocument(
                rfp_package_id=package.id,
                filename=source_path.name,
                storage_path=str(source_path),
                document_type=RfpDocumentType.MAIN_SOLICITATION,
                document_role="original",
                page_count=1,
                extracted_text_md=source_text,
                structure_json={"sections": ["Technical", "Pricing"]},
            )
        )

        proposal = Proposal(
            rfp_package_id=package.id,
            title=TITLE,
            agency="E2E County Treasury",
            naics="522320",
            due_date=date(2031, 8, 31),
            role=ProposalRole.PRIME,
            status=ProposalStatus.AWAITING_COST_BUILD,
            notes="SYNTHETIC E2E ONLY",
            service_line="payment_systems",
            team_approved_at=now,
            evaluation_criteria_json=json.dumps(
                {
                    "evaluation_method": "best_value",
                    "factors": [
                        {
                            "factor_id": "F1",
                            "factor_name": "Payment Solution and Price",
                            "weight_pct": 100,
                            "weight_descriptive": None,
                            "scoring_scale": "Exceptional/Acceptable/Unacceptable",
                            "evidence_required": (
                                "Secure processing, implementation governance, "
                                "and a transparent fee schedule."
                            ),
                            "subfactors": [],
                        }
                    ],
                    "section_l_to_m_map": {
                        "REQ-PAY-001": ["F1"],
                        "REQ-PAY-002": ["F1"],
                    },
                    "trade_off_language": "Technical merit may justify fee tradeoffs.",
                    "lowest_price_clause": None,
                    "extraction_notes": "Synthetic payment workflow fixture.",
                }
            ),
        )
        db.add(proposal)
        db.flush()

        technical = ProposalSection(
            proposal_id=proposal.id,
            section_id="SEC-101",
            section_title="Technical and Implementation Approach",
            section_order=0,
            section_brief=(
                "Describe secure payment processing, implementation governance, "
                "and measurable quality controls."
            ),
            page_limit=5,
            word_limit=1200,
            requires_cost_analysis=False,
            excluded_from_draft=False,
            compliance_items_addressed_json=["REQ-PAY-001"],
            citations_json=[],
            needs_human_placeholders_json=[],
            shortfall_mitigations_applied_json=[],
        )
        cost = ProposalSection(
            proposal_id=proposal.id,
            section_id="SEC-103",
            section_title="Pricing and Fee Narrative",
            section_order=1,
            section_brief=(
                "Present the selected payment-pricing model, proposed rates, "
                "disclosures, and basis of estimate."
            ),
            page_limit=4,
            word_limit=900,
            requires_cost_analysis=True,
            excluded_from_draft=False,
            compliance_items_addressed_json=["REQ-PAY-002"],
            citations_json=[],
            needs_human_placeholders_json=[],
            shortfall_mitigations_applied_json=[],
        )
        db.add_all([technical, cost])
        db.flush()

        db.add_all(
            [
                ComplianceMatrixItem(
                    proposal_id=proposal.id,
                    requirement_id="REQ-PAY-001",
                    requirement_text=(
                        "The contractor shall provide secure cloud modernization "
                        "delivery with transition governance, project management, "
                        "and quality assurance."
                    ),
                    source_doc=source_path.name,
                    source_section="Technical",
                    source_page=1,
                    requirement_type=RequirementType.SHALL,
                    category=RequirementCategory.TECHNICAL,
                    compliance_status=ComplianceStatus.TO_BE_DRAFTED,
                    linked_response_section_id=technical.id,
                    status="active",
                ),
                ComplianceMatrixItem(
                    proposal_id=proposal.id,
                    requirement_id="REQ-PAY-002",
                    requirement_text=(
                        "The offeror shall provide an auditable schedule of all "
                        "proposed transaction and account fees."
                    ),
                    source_doc=source_path.name,
                    source_section="Pricing",
                    source_page=1,
                    requirement_type=RequirementType.SHALL,
                    category=RequirementCategory.PRICING,
                    compliance_status=ComplianceStatus.TO_BE_DRAFTED,
                    linked_response_section_id=cost.id,
                    status="active",
                ),
            ]
        )
        db.add(
            ProposalTeamMember(
                proposal_id=proposal.id,
                role_name="Payment Program Manager",
                person_kind="named",
                assigned_person="Morgan E2E",
                labor_category="Project Manager I",
                wage_band="150k",
                time_allocation_pct=50,
                experience_years=9,
                bio_summary="Synthetic payment implementation lead.",
                phases_active_json=["Implementation"],
                display_order=0,
            )
        )
        db.commit()
        return {"proposal_id": proposal.id}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--cleanup", action="store_true")
    args = parser.parse_args()
    data_dir = _validate_environment()
    payload = _cleanup(data_dir) if args.cleanup else _seed(data_dir)
    print(json.dumps(payload, sort_keys=True))


if __name__ == "__main__":
    main()
