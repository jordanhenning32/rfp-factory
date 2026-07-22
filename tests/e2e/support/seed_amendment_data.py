"""Seed one disposable draft-ready proposal for Amendment browser coverage."""
from __future__ import annotations

import argparse
import json
import os
import sys
from datetime import UTC, date, datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

TITLE = "Synthetic Amendment Lifecycle RFP"


def _validate_environment() -> Path:
    if os.environ.get("APP_ENV", "").strip().lower() != "e2e":
        raise RuntimeError("amendment seed requires APP_ENV=e2e")
    if os.environ.get("RFP_E2E_FAKE_LLM", "") != "1":
        raise RuntimeError("amendment seed requires RFP_E2E_FAKE_LLM=1")

    raw_data_dir = os.environ.get("RFP_DATA_DIR", "").strip()
    if not raw_data_dir:
        raise RuntimeError("amendment seed requires an explicit RFP_DATA_DIR")
    data_dir = Path(raw_data_dir).resolve()
    canonical_data = (PROJECT_ROOT / "data").resolve()
    try:
        data_dir.relative_to(canonical_data)
    except ValueError:
        pass
    else:
        raise RuntimeError(f"refusing to seed canonical data: {data_dir}")

    database_url = os.environ.get("DATABASE_URL", "").replace("\\", "/")
    expected_db = (data_dir / "sqlite.db").as_posix()
    if database_url != f"sqlite:///{expected_db}":
        raise RuntimeError(
            "amendment seed DATABASE_URL must point exactly to sqlite.db "
            f"inside RFP_DATA_DIR; got {database_url!r}"
        )
    if not (data_dir / "sqlite.db").is_file():
        raise RuntimeError("amendment seed requires a migrated E2E database")

    profile_path = data_dir / "company_profile.json"
    try:
        profile = json.loads(profile_path.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:
        raise RuntimeError("amendment seed requires the synthetic profile") from exc
    if not str((profile.get("_meta") or {}).get("version") or "").startswith(
        "e2e-"
    ):
        raise RuntimeError("amendment seed rejected a non-E2E profile")
    return data_dir


def _cleanup() -> dict[str, int | bool]:
    from sqlalchemy import select

    from app.db.session import SessionLocal
    from app.models import Proposal
    from app.services.proposals import delete_proposal

    deleted = 0
    with SessionLocal() as db:
        proposals = db.execute(
            select(Proposal).where(Proposal.title == TITLE)
        ).scalars().all()
        for proposal in proposals:
            result = delete_proposal(db, proposal.id)
            if not result.get("deleted"):
                raise RuntimeError(
                    f"amendment seed cleanup failed for {proposal.id}: {result}"
                )
            deleted += 1
        db.commit()
    return {"deleted": deleted, "ok": True}


def _seed(data_dir: Path) -> dict[str, int]:
    from docx import Document
    from sqlalchemy import select

    from app.core.enums import (
        ComplianceStatus,
        ProposalRole,
        ProposalStatus,
        RequirementCategory,
        RequirementType,
        RfpDocumentType,
    )
    from app.db.session import SessionLocal
    from app.models import (
        ComplianceMatrixItem,
        Proposal,
        ProposalSection,
        RfpPackage,
        RfpPackageDocument,
    )

    now = datetime.now(UTC)
    with SessionLocal() as db:
        existing = db.execute(
            select(Proposal).where(Proposal.title == TITLE)
        ).scalar_one_or_none()
        if existing is not None:
            return {"proposal_id": existing.id, "package_id": existing.rfp_package_id}

        package = RfpPackage(
            uploaded_by="e2e-amendment-seed",
            uploaded_at=now,
            storage_dir="",
            notes="Synthetic package for amendment lifecycle coverage.",
        )
        db.add(package)
        db.flush()

        package_dir = data_dir / "rfp_packages" / str(package.id)
        package_dir.mkdir(parents=True, exist_ok=True)
        package.storage_dir = str(package_dir)
        source_path = package_dir / "synthetic_amendment_base_rfp.docx"
        source = Document()
        source.add_heading("Synthetic Amendment Lifecycle RFP", level=0)
        source.add_paragraph("SYNTHETIC E2E SOLICITATION - NOT A LIVE RFP")
        source.add_heading("Transition Plan", level=1)
        source.add_paragraph(
            "The contractor shall submit the transition plan within 30 "
            "calendar days after award."
        )
        source.add_heading("Legacy Reporting", level=1)
        source.add_paragraph(
            "The contractor must submit a weekly legacy status report."
        )
        source.save(source_path)

        db.add(
            RfpPackageDocument(
                rfp_package_id=package.id,
                filename=source_path.name,
                storage_path=str(source_path),
                document_type=RfpDocumentType.MAIN_SOLICITATION,
                document_role="original",
                page_count=1,
                extracted_text_md=(
                    "The contractor shall submit the transition plan within "
                    "30 calendar days after award. The contractor must submit "
                    "a weekly legacy status report."
                ),
                structure_json={"sections": ["Transition Plan", "Legacy Reporting"]},
            )
        )

        proposal = Proposal(
            rfp_package_id=package.id,
            title=TITLE,
            agency="E2E Department of Amendments",
            naics="541512",
            due_date=date(2032, 4, 30),
            role=ProposalRole.PRIME,
            status=ProposalStatus.DRAFT_READY,
            notes="SYNTHETIC E2E ONLY",
            service_line="it_services",
        )
        db.add(proposal)
        db.flush()

        section = ProposalSection(
            proposal_id=proposal.id,
            section_id="SEC-001",
            section_title="Transition and Reporting Approach",
            section_order=0,
            section_brief="Address transition timing and legacy reporting.",
            page_limit=4,
            word_limit=1000,
            requires_cost_analysis=False,
            excluded_from_draft=False,
            draft_text_markdown=(
                "## Transition and Reporting Approach\n\n"
                "We will submit the transition plan within 30 calendar days "
                "and provide the weekly legacy status report."
            ),
            current_revision_number=1,
            compliance_items_addressed_json=["REQ-001", "REQ-002"],
            citations_json=[],
            needs_human_placeholders_json=[],
            shortfall_mitigations_applied_json=[],
            compliance_drift_pending=False,
        )
        db.add(section)
        db.flush()

        req_1 = ComplianceMatrixItem(
            proposal_id=proposal.id,
            requirement_id="REQ-001",
            requirement_text=(
                "The contractor shall submit the transition plan within 30 "
                "calendar days after award."
            ),
            source_doc=source_path.name,
            source_section="Transition Plan",
            source_page=1,
            requirement_type=RequirementType.SHALL,
            category=RequirementCategory.MANAGEMENT,
            compliance_status=ComplianceStatus.DRAFTED,
            linked_response_section_id=section.id,
            status="active",
        )
        req_2 = ComplianceMatrixItem(
            proposal_id=proposal.id,
            requirement_id="REQ-002",
            requirement_text=(
                "The contractor must submit a weekly legacy status report."
            ),
            source_doc=source_path.name,
            source_section="Legacy Reporting",
            source_page=1,
            requirement_type=RequirementType.MUST,
            category=RequirementCategory.ADMINISTRATIVE,
            compliance_status=ComplianceStatus.DRAFTED,
            linked_response_section_id=section.id,
            status="active",
        )
        db.add_all([req_1, req_2])
        db.commit()
        return {
            "proposal_id": proposal.id,
            "package_id": package.id,
            "section_id": section.id,
        }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--cleanup", action="store_true")
    args = parser.parse_args()
    data_dir = _validate_environment()
    payload = _cleanup() if args.cleanup else _seed(data_dir)
    print(json.dumps(payload, sort_keys=True))


if __name__ == "__main__":
    main()
