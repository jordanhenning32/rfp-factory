"""Seed one isolated proposal at the post-outline team-approval gate.

The downstream browser test intentionally exercises the real product services
from this point forward.  Only the already-completed intake/outline work and a
synthetic MarketScan prerequisite are seeded; team CRUD, cost build, drafting,
reviews, polish, approval, submission, and archive all happen through the UI.
"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
from datetime import UTC, date, datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

TITLE = "Synthetic Downstream IT Services RFP"
PACKAGE_SLUG = "synthetic-downstream-it"


def _validate_environment() -> Path:
    if os.environ.get("APP_ENV", "").strip().lower() != "e2e":
        raise RuntimeError("downstream seed requires APP_ENV=e2e")
    if os.environ.get("RFP_E2E_FAKE_LLM", "") != "1":
        raise RuntimeError("downstream seed requires RFP_E2E_FAKE_LLM=1")

    raw_data_dir = os.environ.get("RFP_DATA_DIR", "").strip()
    if not raw_data_dir:
        raise RuntimeError("downstream seed requires an explicit RFP_DATA_DIR")
    data_dir = Path(raw_data_dir).resolve()
    canonical_data = (PROJECT_ROOT / "data").resolve()
    try:
        data_dir.relative_to(canonical_data)
    except ValueError:
        pass
    else:
        raise RuntimeError(f"refusing to seed canonical data: {data_dir}")

    database_url = os.environ.get("DATABASE_URL", "").replace("\\", "/")
    expected_db = (data_dir / "sqlite.db").as_posix()
    if database_url != f"sqlite:///{expected_db}":
        raise RuntimeError(
            "downstream seed DATABASE_URL must point exactly to sqlite.db "
            f"inside RFP_DATA_DIR; got {database_url!r}"
        )
    if not (data_dir / "sqlite.db").is_file():
        raise RuntimeError("downstream seed requires a migrated E2E database")

    profile_path = data_dir / "company_profile.json"
    try:
        profile = json.loads(profile_path.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:
        raise RuntimeError("downstream seed requires the synthetic profile") from exc
    if not str((profile.get("_meta") or {}).get("version") or "").startswith(
        "e2e-"
    ):
        raise RuntimeError("downstream seed rejected a non-E2E profile")
    return data_dir


def _cleanup(data_dir: Path) -> dict[str, int | bool]:
    from sqlalchemy import select

    from app.db.session import SessionLocal
    from app.models import AgentRun, Proposal, RfpPackage

    deleted = 0
    with SessionLocal() as db:
        proposals = db.execute(
            select(Proposal).where(Proposal.title == TITLE)
        ).scalars().all()
        package_ids = [p.rfp_package_id for p in proposals]
        for proposal in proposals:
            db.query(AgentRun).filter(
                AgentRun.proposal_id == proposal.id
            ).delete(synchronize_session=False)
            db.delete(proposal)
            deleted += 1
        db.flush()
        for package_id in package_ids:
            package = db.get(RfpPackage, package_id)
            if package is not None:
                db.delete(package)
        db.commit()

    package_dir = data_dir / "rfp_packages" / PACKAGE_SLUG
    if package_dir.exists():
        shutil.rmtree(package_dir)
    return {"deleted": deleted, "ok": True}


def _seed(data_dir: Path) -> dict[str, int]:
    from docx import Document
    from sqlalchemy import select

    from app.core.enums import (
        ComplianceStatus,
        ProposalRole,
        ProposalStatus,
        RequirementCategory,
        RequirementType,
        RfpDocumentType,
    )
    from app.db.session import SessionLocal
    from app.models import (
        ComplianceMatrixItem,
        MarketScan,
        Proposal,
        ProposalSection,
        RfpPackage,
        RfpPackageDocument,
    )

    package_dir = data_dir / "rfp_packages" / PACKAGE_SLUG
    source_path = package_dir / "synthetic_downstream_rfp.docx"
    package_dir.mkdir(parents=True, exist_ok=True)
    source_text = (
        "SYNTHETIC E2E SOLICITATION. The contractor shall design and "
        "implement a secure cloud modernization approach. The offeror must "
        "provide transition governance, project management, and quality "
        "assurance. Period of performance is 12 months under a firm-fixed-"
        "price contract."
    )
    document = Document()
    document.add_heading("Synthetic Downstream IT Services RFP", level=0)
    document.add_paragraph("SYNTHETIC E2E SOLICITATION — NOT A LIVE RFP")
    document.add_heading("Technical and Management Requirements", level=1)
    document.add_paragraph(source_text)
    document.add_heading("Cost Volume", level=1)
    document.add_paragraph("The offeror shall provide a complete cost volume.")
    document.save(source_path)

    now = datetime.now(UTC)
    with SessionLocal() as db:
        existing = db.execute(
            select(Proposal).where(Proposal.title == TITLE)
        ).scalar_one_or_none()
        if existing is not None:
            return {"proposal_id": existing.id}

        package = RfpPackage(
            uploaded_by="e2e-downstream-seed",
            uploaded_at=now,
            storage_dir=str(package_dir),
            notes="Synthetic package for downstream behavioral coverage.",
        )
        db.add(package)
        db.flush()
        db.add(
            RfpPackageDocument(
                rfp_package_id=package.id,
                filename=source_path.name,
                storage_path=str(source_path),
                document_type=RfpDocumentType.MAIN_SOLICITATION,
                document_role="original",
                page_count=1,
                extracted_text_md=source_text,
                structure_json={"sections": ["Technical", "Management", "Cost"]},
            )
        )

        proposal = Proposal(
            rfp_package_id=package.id,
            title=TITLE,
            agency="E2E Department of Delivery",
            naics="541512",
            due_date=date(2031, 6, 30),
            role=ProposalRole.PRIME,
            status=ProposalStatus.AWAITING_TEAM_APPROVAL,
            notes="SYNTHETIC E2E ONLY",
            service_line="it_services",
            proposed_scenario="MEDIUM",
            evaluation_criteria_json=json.dumps(
                {
                    "evaluation_method": "best_value",
                    "factors": [
                        {
                            "factor_id": "F1",
                            "factor_name": "Technical and Management Approach",
                            "weight_pct": 70,
                            "weight_descriptive": None,
                            "scoring_scale": "Exceptional/Acceptable/Unacceptable",
                            "evidence_required": (
                                "Secure cloud delivery, transition governance, "
                                "project management, and quality assurance."
                            ),
                            "subfactors": [],
                        }
                    ],
                    "section_l_to_m_map": {
                        "REQ-DOWN-001": ["F1"],
                        "REQ-DOWN-002": ["F1"],
                    },
                    "trade_off_language": "Technical merit may justify price tradeoffs.",
                    "lowest_price_clause": None,
                    "extraction_notes": "Synthetic downstream fixture.",
                }
            ),
        )
        db.add(proposal)
        db.flush()

        sections = [
            ProposalSection(
                proposal_id=proposal.id,
                section_id="SEC-101",
                section_title="Technical Approach",
                section_order=0,
                section_brief=(
                    "Describe secure cloud modernization delivery and a testable "
                    "implementation approach."
                ),
                page_limit=5,
                word_limit=1200,
                requires_cost_analysis=False,
                excluded_from_draft=False,
                compliance_items_addressed_json=["REQ-DOWN-001"],
                citations_json=[],
                needs_human_placeholders_json=[],
                shortfall_mitigations_applied_json=[],
            ),
            ProposalSection(
                proposal_id=proposal.id,
                section_id="SEC-102",
                section_title="Management and Quality Approach",
                section_order=1,
                section_brief=(
                    "Describe transition governance, project management, and "
                    "quality assurance controls."
                ),
                page_limit=3,
                word_limit=800,
                requires_cost_analysis=False,
                excluded_from_draft=False,
                compliance_items_addressed_json=["REQ-DOWN-002"],
                citations_json=[],
                needs_human_placeholders_json=[],
                shortfall_mitigations_applied_json=[],
            ),
            ProposalSection(
                proposal_id=proposal.id,
                section_id="SEC-103",
                section_title="Cost Volume",
                section_order=2,
                section_brief="Present the selected scenario and its cost basis.",
                page_limit=3,
                word_limit=700,
                requires_cost_analysis=True,
                excluded_from_draft=False,
                compliance_items_addressed_json=["REQ-DOWN-003"],
                citations_json=[],
                needs_human_placeholders_json=[],
                shortfall_mitigations_applied_json=[],
            ),
        ]
        db.add_all(sections)
        db.flush()

        requirements = [
            ComplianceMatrixItem(
                proposal_id=proposal.id,
                requirement_id="REQ-DOWN-001",
                requirement_text=(
                    "The contractor shall design and implement a secure cloud "
                    "modernization approach."
                ),
                source_doc=source_path.name,
                source_section="Technical",
                source_page=1,
                requirement_type=RequirementType.SHALL,
                category=RequirementCategory.TECHNICAL,
                compliance_status=ComplianceStatus.TO_BE_DRAFTED,
                linked_response_section_id=sections[0].id,
                status="active",
            ),
            ComplianceMatrixItem(
                proposal_id=proposal.id,
                requirement_id="REQ-DOWN-002",
                requirement_text=(
                    "The offeror must provide transition governance, project "
                    "management, and quality assurance."
                ),
                source_doc=source_path.name,
                source_section="Management",
                source_page=1,
                requirement_type=RequirementType.MUST,
                category=RequirementCategory.MANAGEMENT,
                compliance_status=ComplianceStatus.TO_BE_DRAFTED,
                linked_response_section_id=sections[1].id,
                status="active",
            ),
            ComplianceMatrixItem(
                proposal_id=proposal.id,
                requirement_id="REQ-DOWN-003",
                requirement_text="The offeror shall provide a complete cost volume.",
                source_doc=source_path.name,
                source_section="Cost",
                source_page=1,
                requirement_type=RequirementType.SHALL,
                category=RequirementCategory.PRICING,
                compliance_status=ComplianceStatus.TO_BE_DRAFTED,
                linked_response_section_id=sections[2].id,
                status="active",
            ),
        ]
        db.add_all(requirements)

        db.add(
            MarketScan(
                proposal_id=proposal.id,
                market_band_low_usd=175000,
                market_band_mid_usd=260000,
                market_band_high_usd=375000,
                methodology=(
                    "Synthetic E2E prerequisite seam: bounded market band for "
                    "deterministic downstream pricing behavior."
                ),
            )
        )
        db.commit()
        return {"proposal_id": proposal.id}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--cleanup", action="store_true")
    args = parser.parse_args()
    data_dir = _validate_environment()
    payload = _cleanup(data_dir) if args.cleanup else _seed(data_dir)
    print(json.dumps(payload, sort_keys=True))


if __name__ == "__main__":
    main()
