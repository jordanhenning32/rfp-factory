"""Seed one isolated IT-services proposal at the market-research gate.

Intake, outline creation, and team approval are deliberate fixture seams. The
browser test starts with no MarketScan and drives the production dual-provider
market-research job through the real Cost-tab control.
"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
from datetime import UTC, date, datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

TITLE = "Synthetic IT Market Research RFP"
PACKAGE_SLUG = "synthetic-it-market-research"
UPLOADED_BY = "e2e-it-market-seed"


def _validate_environment() -> Path:
    if os.environ.get("APP_ENV", "").strip().lower() != "e2e":
        raise RuntimeError("IT market seed requires APP_ENV=e2e")
    if os.environ.get("RFP_E2E_FAKE_LLM", "") != "1":
        raise RuntimeError("IT market seed requires RFP_E2E_FAKE_LLM=1")

    raw_data_dir = os.environ.get("RFP_DATA_DIR", "").strip()
    if not raw_data_dir:
        raise RuntimeError("IT market seed requires an explicit RFP_DATA_DIR")
    data_dir = Path(raw_data_dir).resolve()
    canonical_data = (PROJECT_ROOT / "data").resolve()
    try:
        data_dir.relative_to(canonical_data)
    except ValueError:
        pass
    else:
        raise RuntimeError(f"refusing to seed canonical data: {data_dir}")

    expected_db = (data_dir / "sqlite.db").as_posix()
    database_url = os.environ.get("DATABASE_URL", "").replace("\\", "/")
    if database_url != f"sqlite:///{expected_db}":
        raise RuntimeError(
            "IT market seed DATABASE_URL must point exactly to sqlite.db "
            f"inside RFP_DATA_DIR; got {database_url!r}"
        )
    if not (data_dir / "sqlite.db").is_file():
        raise RuntimeError("IT market seed requires a migrated E2E database")

    profile_path = data_dir / "company_profile.json"
    try:
        profile = json.loads(profile_path.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:
        raise RuntimeError("IT market seed requires the synthetic profile") from exc
    if not str((profile.get("_meta") or {}).get("version") or "").startswith(
        "e2e-"
    ):
        raise RuntimeError("IT market seed rejected a non-E2E profile")
    return data_dir


def _cleanup(data_dir: Path) -> dict[str, int | bool]:
    from sqlalchemy import select

    from app.db.session import SessionLocal
    from app.models import AgentRun, Proposal, RfpPackage

    package_dir = (data_dir / "rfp_packages" / PACKAGE_SLUG).resolve()
    deleted = 0
    with SessionLocal() as db:
        proposals = db.execute(
            select(Proposal).where(Proposal.title == TITLE)
        ).scalars().all()
        packages: list[RfpPackage] = []
        for proposal in proposals:
            package = db.get(RfpPackage, proposal.rfp_package_id)
            if package is None:
                raise RuntimeError(
                    f"refusing unsafe cleanup: proposal {proposal.id} has no package"
                )
            if (
                package.uploaded_by != UPLOADED_BY
                or Path(package.storage_dir).resolve() != package_dir
            ):
                raise RuntimeError(
                    "refusing unsafe cleanup: matching title is not owned by "
                    "the isolated IT market seed"
                )
            packages.append(package)
            db.query(AgentRun).filter(
                AgentRun.proposal_id == proposal.id
            ).delete(synchronize_session=False)
            db.delete(proposal)
            deleted += 1
        db.flush()
        for package in packages:
            db.delete(package)
        db.commit()

    if package_dir.exists():
        shutil.rmtree(package_dir)
    return {"deleted": deleted, "ok": True}


def _seed(data_dir: Path) -> dict[str, int]:
    from docx import Document
    from sqlalchemy import func, select

    from app.core.enums import (
        ComplianceStatus,
        ProposalRole,
        ProposalStatus,
        RequirementCategory,
        RequirementType,
        RfpDocumentType,
    )
    from app.db.session import SessionLocal
    from app.models import (
        ComplianceMatrixItem,
        MarketScan,
        Proposal,
        ProposalSection,
        ProposalTeamMember,
        RfpPackage,
        RfpPackageDocument,
    )

    # Remove a stale fixture left by an interrupted local run. The cleanup is
    # guarded by the seed-specific uploader and exact disposable path.
    _cleanup(data_dir)
    package_dir = data_dir / "rfp_packages" / PACKAGE_SLUG
    source_path = package_dir / "synthetic_it_market_rfp.docx"
    package_dir.mkdir(parents=True, exist_ok=True)
    source_text = (
        "SYNTHETIC E2E SOLICITATION — NOT A LIVE RFP. The contractor shall "
        "modernize a browser-based records platform, migrate workloads to a "
        "secure cloud environment, and provide delivery governance. The period "
        "of performance is 12 months. The offeror shall propose qualified "
        "technical and program-management personnel."
    )
    document = Document()
    document.add_heading(TITLE, level=0)
    document.add_paragraph("SYNTHETIC E2E SOLICITATION — NOT A LIVE RFP")
    document.add_heading("Technical and Management Requirements", level=1)
    document.add_paragraph(source_text)
    document.add_heading("Cost Volume", level=1)
    document.add_paragraph("The offeror shall provide a complete cost volume.")
    document.save(source_path)

    now = datetime.now(UTC)
    with SessionLocal() as db:
        package = RfpPackage(
            uploaded_by=UPLOADED_BY,
            uploaded_at=now,
            storage_dir=str(package_dir),
            notes="Synthetic package for isolated IT market research coverage.",
        )
        db.add(package)
        db.flush()
        db.add(
            RfpPackageDocument(
                rfp_package_id=package.id,
                filename=source_path.name,
                storage_path=str(source_path),
                document_type=RfpDocumentType.MAIN_SOLICITATION,
                document_role="original",
                page_count=1,
                extracted_text_md=source_text,
                structure_json={
                    "sections": ["Technical", "Management", "Personnel", "Cost"]
                },
            )
        )

        proposal = Proposal(
            rfp_package_id=package.id,
            title=TITLE,
            agency="E2E Digital Services Agency",
            naics="541512",
            due_date=date(2031, 9, 30),
            role=ProposalRole.PRIME,
            status=ProposalStatus.AWAITING_COST_BUILD,
            notes="SYNTHETIC E2E ONLY",
            service_line="it_services",
            team_approved_at=now,
            proposed_scenario="MEDIUM",
            evaluation_criteria_json=json.dumps(
                {
                    "evaluation_method": "best_value",
                    "factors": [
                        {
                            "factor_id": "F1",
                            "factor_name": "Technical and Management Approach",
                            "weight_pct": 70,
                            "weight_descriptive": None,
                            "scoring_scale": "Exceptional/Acceptable/Unacceptable",
                            "evidence_required": (
                                "Secure cloud modernization, transition governance, "
                                "and qualified personnel."
                            ),
                            "subfactors": [],
                        }
                    ],
                    "section_l_to_m_map": {
                        "REQ-MKT-001": ["F1"],
                        "REQ-MKT-002": ["F1"],
                        "REQ-MKT-003": ["F1"],
                    },
                    "trade_off_language": (
                        "Technical merit may justify a price tradeoff."
                    ),
                    "lowest_price_clause": None,
                    "extraction_notes": "Synthetic IT market fixture.",
                }
            ),
        )
        db.add(proposal)
        db.flush()

        technical = ProposalSection(
            proposal_id=proposal.id,
            section_id="SEC-MKT-101",
            section_title="Technical and Management Approach",
            section_order=0,
            section_brief=(
                "Describe secure modernization delivery, governance, and the "
                "approved team."
            ),
            page_limit=6,
            word_limit=1500,
            requires_cost_analysis=False,
            excluded_from_draft=False,
            compliance_items_addressed_json=[
                "REQ-MKT-001",
                "REQ-MKT-002",
                "REQ-MKT-003",
            ],
            citations_json=[],
            needs_human_placeholders_json=[],
            shortfall_mitigations_applied_json=[],
        )
        cost = ProposalSection(
            proposal_id=proposal.id,
            section_id="SEC-MKT-102",
            section_title="Cost Volume",
            section_order=1,
            section_brief="Present the selected scenario and basis of estimate.",
            page_limit=3,
            word_limit=700,
            requires_cost_analysis=True,
            excluded_from_draft=False,
            compliance_items_addressed_json=["REQ-MKT-004"],
            citations_json=[],
            needs_human_placeholders_json=[],
            shortfall_mitigations_applied_json=[],
        )
        db.add_all([technical, cost])
        db.flush()

        requirements = [
            (
                "REQ-MKT-001",
                "The contractor shall modernize the browser-based records platform "
                "during a 12-month period of performance.",
                "Technical",
                RequirementCategory.TECHNICAL,
                technical.id,
            ),
            (
                "REQ-MKT-002",
                "The offeror must provide secure cloud migration and delivery "
                "governance.",
                "Management",
                RequirementCategory.MANAGEMENT,
                technical.id,
            ),
            (
                "REQ-MKT-003",
                "The offeror shall propose qualified technical and program-management "
                "personnel.",
                "Personnel",
                RequirementCategory.PERSONNEL,
                technical.id,
            ),
            (
                "REQ-MKT-004",
                "The offeror shall provide a complete cost volume.",
                "Cost",
                RequirementCategory.PRICING,
                cost.id,
            ),
        ]
        for requirement_id, text, section, category, linked_id in requirements:
            db.add(
                ComplianceMatrixItem(
                    proposal_id=proposal.id,
                    requirement_id=requirement_id,
                    requirement_text=text,
                    source_doc=source_path.name,
                    source_section=section,
                    source_page=1,
                    requirement_type=(
                        RequirementType.MUST
                        if " must " in f" {text.lower()} "
                        else RequirementType.SHALL
                    ),
                    category=category,
                    compliance_status=ComplianceStatus.TO_BE_DRAFTED,
                    linked_response_section_id=linked_id,
                    status="active",
                )
            )

        db.add(
            ProposalTeamMember(
                proposal_id=proposal.id,
                role_name="Modernization Program Manager",
                person_kind="named",
                assigned_person="Morgan E2E",
                labor_category="Project Manager I",
                wage_band="150k",
                time_allocation_pct=50,
                experience_years=9,
                bio_summary="Synthetic approved modernization delivery lead.",
                phases_active_json=["Delivery"],
                display_order=0,
            )
        )
        db.commit()

        scan_count = db.execute(
            select(func.count(MarketScan.id)).where(
                MarketScan.proposal_id == proposal.id
            )
        ).scalar_one()
        if scan_count != 0:
            raise RuntimeError("IT market seed unexpectedly created a MarketScan")
        return {"proposal_id": proposal.id, "package_id": package.id}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--cleanup", action="store_true")
    args = parser.parse_args()
    data_dir = _validate_environment()
    payload = _cleanup(data_dir) if args.cleanup else _seed(data_dir)
    print(json.dumps(payload, sort_keys=True))


if __name__ == "__main__":
    main()
