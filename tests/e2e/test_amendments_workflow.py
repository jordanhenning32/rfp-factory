"""Real-browser Amendment upload, delta-apply, and duplicate-safety coverage."""
from __future__ import annotations

import json
import re
import sqlite3
import subprocess
import sys
import time
from collections.abc import Callable
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from playwright.sync_api import expect

from tests.e2e.conftest import BrowserSession, E2EServer

pytestmark = pytest.mark.e2e

PROJECT_ROOT = Path(__file__).resolve().parents[2]
TITLE = "Synthetic Amendment Lifecycle RFP"
AMENDMENT_FILENAME = "synthetic_amendment_0002.docx"


@pytest.fixture()
def amendment_seed(e2e_server: E2EServer):
    ledger_path = e2e_server.workspace.artifacts / "llm_calls.jsonl"
    ledger_start = (
        len(ledger_path.read_text(encoding="utf-8").splitlines())
        if ledger_path.exists()
        else 0
    )
    script = (
        PROJECT_ROOT / "tests" / "e2e" / "support" / "seed_amendment_data.py"
    )
    result = subprocess.run(
        [sys.executable, str(script)],
        cwd=PROJECT_ROOT,
        env=e2e_server.workspace.environment,
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )
    seed_log = e2e_server.workspace.artifacts / "amendment_seed.log"
    seed_log.write_text(
        (result.stdout or "") + (result.stderr or ""), encoding="utf-8"
    )
    if result.returncode != 0:
        pytest.fail(
            f"synthetic amendment seed failed (exit {result.returncode}); "
            f"see {seed_log}",
            pytrace=False,
        )
    try:
        payload = json.loads((result.stdout or "").strip().splitlines()[-1])
    except (IndexError, ValueError) as exc:
        pytest.fail(
            f"amendment seed returned invalid JSON: {exc}; see {seed_log}",
            pytrace=False,
        )
    payload["ledger_start"] = ledger_start
    yield payload

    cleanup = subprocess.run(
        [sys.executable, str(script), "--cleanup"],
        cwd=PROJECT_ROOT,
        env=e2e_server.workspace.environment,
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )
    cleanup_log = e2e_server.workspace.artifacts / "amendment_cleanup.log"
    cleanup_log.write_text(
        (cleanup.stdout or "") + (cleanup.stderr or ""), encoding="utf-8"
    )
    if cleanup.returncode != 0:
        pytest.fail(
            f"synthetic amendment cleanup failed (exit {cleanup.returncode}); "
            f"see {cleanup_log}",
            pytrace=False,
        )


@pytest.fixture()
def amendment_browser(
    amendment_seed: dict[str, int],
    browser_session: BrowserSession,
) -> tuple[BrowserSession, dict[str, int]]:
    return browser_session, amendment_seed


def _fetchone(
    database_path: Path,
    sql: str,
    params: tuple[Any, ...] = (),
) -> sqlite3.Row | None:
    with sqlite3.connect(database_path, timeout=2.0) as db:
        db.row_factory = sqlite3.Row
        return db.execute(sql, params).fetchone()


def _fetchall(
    database_path: Path,
    sql: str,
    params: tuple[Any, ...] = (),
) -> list[sqlite3.Row]:
    with sqlite3.connect(database_path, timeout=2.0) as db:
        db.row_factory = sqlite3.Row
        return db.execute(sql, params).fetchall()


def _eventually(
    read: Callable[[], Any],
    predicate: Callable[[Any], bool],
    *,
    description: str,
    timeout: float = 20.0,
) -> Any:
    deadline = time.monotonic() + timeout
    last = None
    while time.monotonic() < deadline:
        last = read()
        if predicate(last):
            return last
        time.sleep(0.05)
    raise AssertionError(f"timed out waiting for {description}; last={last!r}")


def _proposal_tab(page, label: str):
    label_node = page.locator(".q-tab__label").filter(
        has_text=re.compile(rf"^\s*{re.escape(label)}\s*$", re.I)
    )
    return page.get_by_role("tab").filter(has=label_node).first


def _open_amendments_tab(page) -> None:
    tab = _proposal_tab(page, "Amendments & Q&A")
    tab.wait_for(state="visible", timeout=10_000)
    tab.click()
    page.get_by_text("Upload Amendment", exact=True).wait_for(
        state="visible", timeout=10_000
    )
    expect(tab).to_have_attribute("aria-selected", "true", timeout=10_000)


def _generate_amendment(path: Path) -> bytes:
    document = Document()
    document.add_heading("Amendment 0002", level=0)
    document.add_paragraph("SYNTHETIC E2E AMENDMENT - NOT A LIVE PROCUREMENT")
    document.add_heading("Transition Plan Change", level=1)
    document.add_paragraph(
        "The contractor shall submit the transition plan within 15 calendar "
        "days after award."
    )
    document.add_heading("Legacy Reporting Removal", level=1)
    document.add_paragraph(
        "Amendment 0002 removes the weekly legacy status report requirement."
    )
    document.add_heading("New Monthly Reporting", level=1)
    document.add_paragraph(
        "The contractor shall submit a monthly amendment impact report."
    )
    document.save(path)
    return path.read_bytes()


def _stage_amendment(page, amendment_path: Path, *, sequence: int) -> None:
    card = page.locator(".q-card").filter(has_text="Upload Amendment").first
    card.locator('input[type="number"]').fill(str(sequence))
    upload = card.locator('input[type="file"]')
    assert upload.count() == 1
    upload.set_input_files(str(amendment_path))
    page.get_by_text(
        re.compile(rf"^Staged amendment {re.escape(AMENDMENT_FILENAME)} "),
    ).wait_for(state="visible", timeout=10_000)
    card.get_by_role("button", name="Upload amendment", exact=True).click()


def _new_ledger_calls(path: Path, start: int) -> list[tuple[str, str, str | None]]:
    entries = [
        json.loads(line)
        for line in path.read_text(encoding="utf-8").splitlines()[start:]
        if line.strip()
    ]
    return [
        (entry["method"], entry["agent_name"], entry.get("tool_name"))
        for entry in entries
    ]


def test_amendment_docx_delta_apply_timeline_and_duplicate_skip(
    amendment_browser: tuple[BrowserSession, dict[str, int]],
    tmp_path: Path,
) -> None:
    session, seed = amendment_browser
    page = session.page
    server = session.server
    database_path = server.workspace.database_path
    proposal_id = int(seed["proposal_id"])
    package_id = int(seed["package_id"])
    amendment_path = tmp_path / AMENDMENT_FILENAME
    amendment_bytes = _generate_amendment(amendment_path)

    response = page.goto(
        f"{server.base_url}/proposals/{proposal_id}",
        wait_until="load",
    )
    assert response is not None and response.status == 200
    page.get_by_text(TITLE, exact=True).wait_for(state="visible", timeout=15_000)
    _open_amendments_tab(page)
    page.get_by_text("No amendments uploaded yet.", exact=False).wait_for()

    # A successful upload intentionally schedules a product-side reload after
    # 500 ms. Synchronize on that navigation so a later explicit refresh can
    # never race it and produce Chromium's net::ERR_ABORTED.
    with page.expect_navigation(wait_until="load", timeout=10_000):
        _stage_amendment(page, amendment_path, sequence=2)

    run = _eventually(
        lambda: _fetchone(
            database_path,
            """
            SELECT ar.id AS run_id, ar.document_id, ar.status, ar.report_json,
                   ar.error_text, d.filename, d.storage_path, d.document_role,
                   d.sequence_number, d.structure_json
            FROM amendment_runs AS ar
            JOIN rfp_package_documents AS d ON d.id = ar.document_id
            WHERE ar.proposal_id = ?
            ORDER BY ar.id DESC LIMIT 1
            """,
            (proposal_id,),
        ),
        lambda row: row is not None and row["status"] in {"completed", "failed"},
        description="terminal amendment ingestion",
    )
    assert run["status"] == "completed", run["error_text"]
    assert run["filename"] == AMENDMENT_FILENAME
    assert run["document_role"] == "amendment"
    assert run["sequence_number"] == 2
    stored_path = Path(run["storage_path"]).resolve()
    expected_package_dir = (
        server.workspace.root / "rfp_packages" / str(package_id)
    ).resolve()
    assert stored_path.parent == expected_package_dir
    assert stored_path.read_bytes() == amendment_bytes
    assert json.loads(run["structure_json"])["content_sha256"]

    report = json.loads(run["report_json"])
    assert report == {
        "n_new": 1,
        "n_modified": 1,
        "n_removed": 1,
        "sections_marked_stale": ["SEC-001"],
        "due_date_changed": False,
        "page_limit_changes": [],
    }

    requirement_rows = _fetchall(
        database_path,
        """
        SELECT id, requirement_id, requirement_text, source_doc,
               amendment_origin, status, superseded_by_id,
               compliance_status, linked_response_section_id
        FROM compliance_matrix_items
        WHERE proposal_id = ?
        ORDER BY id
        """,
        (proposal_id,),
    )
    assert len(requirement_rows) == 4
    old_req_1 = next(
        row
        for row in requirement_rows
        if row["requirement_id"] == "REQ-001" and row["status"] == "superseded"
    )
    current_req_1 = next(
        row
        for row in requirement_rows
        if row["requirement_id"] == "REQ-001" and row["status"] == "active"
    )
    removed_req_2 = next(
        row for row in requirement_rows if row["requirement_id"] == "REQ-002"
    )
    new_req_3 = next(
        row for row in requirement_rows if row["requirement_id"] == "REQ-003"
    )
    assert old_req_1["superseded_by_id"] == current_req_1["id"]
    assert current_req_1["requirement_text"] == (
        "The contractor shall submit the transition plan within 15 calendar "
        "days after award."
    )
    assert current_req_1["source_doc"] == AMENDMENT_FILENAME
    assert current_req_1["amendment_origin"] == AMENDMENT_FILENAME
    assert current_req_1["compliance_status"] == "to_be_drafted"
    assert current_req_1["linked_response_section_id"] == int(seed["section_id"])
    assert removed_req_2["status"] == "removed"
    assert removed_req_2["amendment_origin"] == AMENDMENT_FILENAME
    assert new_req_3["status"] == "active"
    assert new_req_3["source_doc"] == AMENDMENT_FILENAME
    assert new_req_3["amendment_origin"] == AMENDMENT_FILENAME
    assert new_req_3["requirement_text"] == (
        "The contractor shall submit a monthly amendment impact report."
    )

    section = _fetchone(
        database_path,
        "SELECT compliance_drift_pending, draft_text_markdown "
        "FROM proposal_sections WHERE proposal_id = ? AND section_id = 'SEC-001'",
        (proposal_id,),
    )
    assert section is not None
    assert section["compliance_drift_pending"] == 1
    assert "30 calendar days" in section["draft_text_markdown"]
    proposal = _fetchone(
        database_path,
        "SELECT status FROM proposals WHERE id = ?",
        (proposal_id,),
    )
    assert proposal is not None and proposal["status"] == "draft_ready"
    delta_run = _fetchone(
        database_path,
        "SELECT status FROM agent_runs WHERE proposal_id = ? "
        "AND agent_name = 'compliance_matrix_delta'",
        (proposal_id,),
    )
    assert delta_run is not None and delta_run["status"] == "completed"

    # Reload the real product route after the daemon completes so the latest
    # report and timeline are rendered from persistence, not transient state.
    page.goto(
        f"{server.base_url}/proposals/{proposal_id}",
        wait_until="load",
    )
    _open_amendments_tab(page)
    for text in (
        "new: 1",
        "modified: 1",
        "removed: 1",
        "sections flagged: 1",
        "SEC-001",
        AMENDMENT_FILENAME,
        "seq #2",
        "completed",
    ):
        page.get_by_text(text, exact=True).first.wait_for(
            state="visible", timeout=10_000
        )

    # Stage the same bytes again. The browser-visible dedup dialog must offer
    # the safe path, and choosing it must create neither a document nor a run.
    docs_before = _fetchone(
        database_path,
        "SELECT COUNT(*) AS n FROM rfp_package_documents "
        "WHERE rfp_package_id = ?",
        (package_id,),
    )["n"]
    runs_before = _fetchone(
        database_path,
        "SELECT COUNT(*) AS n FROM amendment_runs WHERE proposal_id = ?",
        (proposal_id,),
    )["n"]
    _stage_amendment(page, amendment_path, sequence=2)
    dialog = page.get_by_role("dialog").last
    dialog.get_by_text("match existing", exact=False).wait_for(
        state="visible", timeout=10_000
    )
    dialog.get_by_role("button", name="Skip duplicates", exact=True).click()
    page.get_by_text(
        "All staged files matched existing uploads — skipped.", exact=True
    ).wait_for(state="visible", timeout=10_000)
    page.wait_for_timeout(300)
    assert _fetchone(
        database_path,
        "SELECT COUNT(*) AS n FROM rfp_package_documents "
        "WHERE rfp_package_id = ?",
        (package_id,),
    )["n"] == docs_before
    assert _fetchone(
        database_path,
        "SELECT COUNT(*) AS n FROM amendment_runs WHERE proposal_id = ?",
        (proposal_id,),
    )["n"] == runs_before

    package_files = sorted(path.name for path in expected_package_dir.iterdir())
    assert package_files == [
        AMENDMENT_FILENAME,
        "synthetic_amendment_base_rfp.docx",
    ]
    ledger_path = server.workspace.artifacts / "llm_calls.jsonl"
    assert _new_ledger_calls(ledger_path, int(seed["ledger_start"])) == [
        ("call_tool", "compliance_matrix_delta", "report_compliance_delta")
    ]
