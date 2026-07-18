"""Document text extraction (PDF + DOCX + XLSX + plain text).

PDF: pdfplumber, native text only. OCR fallback (Step 4) lands later.
DOCX: python-docx, paragraphs + table cells.
XLSX: openpyxl, all sheets flattened to a pipe-delimited text view.
TXT/MD/CSV: direct UTF-8 read.

All extractors share the same return shape: (text, page_count). page_count
is 0 for non-paginated formats (DOCX/TXT/XLSX).
"""

from __future__ import annotations

import logging
from pathlib import Path

import pdfplumber

log = logging.getLogger(__name__)

DEFAULT_MAX_PDF_PAGES = 300
DEFAULT_MAX_TEXT_CHARS = 500_000


def _cap_text(text: str, max_chars: int | None = DEFAULT_MAX_TEXT_CHARS) -> str:
    if max_chars is None or len(text) <= max_chars:
        return text
    return text[:max_chars] + f"\n\n... (truncated at {max_chars} characters)"


def extract_pdf_text(path: Path | str, max_pages: int | None = DEFAULT_MAX_PDF_PAGES) -> tuple[str, int]:
    """Return (concatenated text with page markers, total page count).

    `max_pages` caps how many pages are read — useful for cheap metadata
    passes that only need the cover/SOW summary. None = read everything.
    """
    path = Path(path)
    pages_text: list[str] = []
    total = 0
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            if max_pages is not None and i >= max_pages:
                break
            try:
                text = page.extract_text() or ""
            except Exception:
                log.exception("pdfplumber extract_text failed on page %d of %s", i + 1, path.name)
                text = ""
            pages_text.append(f"--- Page {i + 1} ---\n{text}")
    return _cap_text("\n\n".join(pages_text)), total


def extract_docx_text(path: Path | str) -> tuple[str, int]:
    """Pull paragraphs + table cell text from a .docx file. Lazy-imports
    python-docx so the module can be used in environments where DOCX
    support isn't installed."""
    from docx import Document  # local import — keeps cold start fast

    path = Path(path)
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Tables — flatten cells row by row.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _cap_text("\n".join(parts)), 0


def extract_xlsx_text(
    path: Path | str,
    *,
    max_rows_per_sheet: int = 500,
    max_cells_per_row: int = 50,
) -> tuple[str, int]:
    """Flatten every sheet of an xlsx file to a pipe-delimited text view.

    Caps applied for cost control: max_rows_per_sheet limits the row count
    per sheet (rare for KB docs to need more), max_cells_per_row trims very
    wide sheets. Returns (text, sheet_count). page_count slot is reused for
    sheet count.
    """
    from openpyxl import load_workbook  # local import keeps cold start fast

    path = Path(path)
    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    parts: list[str] = []
    sheet_count = 0
    try:
        for sheet_name in wb.sheetnames:
            sheet_count += 1
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= max_rows_per_sheet:
                    parts.append(f"... (truncated at {max_rows_per_sheet} rows)")
                    break
                cells = [str(c) if c is not None else "" for c in row[:max_cells_per_row]]
                # Skip fully-empty rows so the output isn't padded with blanks.
                if any(cell.strip() for cell in cells):
                    parts.append(" | ".join(cells))
                row_count += 1
    finally:
        wb.close()
    return _cap_text("\n".join(parts)), sheet_count


def extract_text_for_path(path: Path | str) -> tuple[str, int]:
    """Dispatcher: pick the right extractor based on file suffix.

    Returns ("", 0) for unsupported types — caller can decide whether
    to warn or fall through.
    """
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".xlsx":
        return extract_xlsx_text(path)
    if suffix in {".txt", ".md", ".markdown", ".csv"}:
        try:
            with path.open("r", encoding="utf-8", errors="replace") as fh:
                return _cap_text(fh.read(DEFAULT_MAX_TEXT_CHARS + 1)), 0
        except Exception:
            log.exception("failed to read text file %s", path)
            return "", 0
    log.warning("extract_text_for_path: unsupported file type %s for %s", suffix, path.name)
    return "", 0
