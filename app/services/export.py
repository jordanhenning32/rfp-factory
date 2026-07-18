"""Proposal export — compile the drafted proposal into a downloadable
DOCX file.

This is the canonical place for "render the proposal as a deliverable
artifact." Multiple UI surfaces (Final Polish preview modal, eventual
Submission Checklist export, etc.) call into the same compiler so the
output is identical regardless of which button the user clicks.

Conversion strategy: walk the compiled markdown line-by-line and emit
docx paragraphs / runs / tables. Covers the markdown subset our agents
actually produce:

  - `## SEC-### — Title` → Heading 1
  - `### Sub` / `#### Sub` → Heading 2 / Heading 3
  - `**bold**` / `*italic*` / `` `code` `` runs
  - `- item` / `* item` bullet lists
  - `1. item` numbered lists
  - `| col | col |` markdown tables (with alignment row)
  - `> quote` block-quote (rendered as indented italic paragraph)
  - blank line separates paragraphs

Anything we don't recognize falls through as plain text. Federal
evaluators read final DOCX in Word — minor formatting drift is fine
as long as the content is correct.

Page setup: standard 1" margins, Calibri 11pt body, Calibri 14/13/12
for H1/H2/H3. The user can restyle in Word; we just want the
information density and structure to be reviewable as-is.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from app.services.sections import compile_proposal_markdown
from app.services.submission_commitments import (
    get_submission_checklist_snapshot,
)

log = logging.getLogger(__name__)


# ---- Inline-format walker -------------------------------------------------

_INLINE_RE = re.compile(
    r"(\*\*[^*\n]+?\*\*"  # **bold**
    r"|\*[^*\n]+?\*"  # *italic*
    r"|__[^_\n]+?__"  # __bold__
    r"|_[^_\n]+?_"  # _italic_
    r"|`[^`\n]+?`"  # `code`
    r")",
    re.UNICODE,
)


def _emit_inline_runs(paragraph, text: str) -> None:
    """Walk a single line of markdown text, emitting docx runs for
    bold / italic / code / plain spans. Nested formatting (e.g.,
    bold-inside-italic) is NOT supported — markdown the writer
    actually produces rarely needs it; we'd need a real parser to
    handle nesting cleanly."""
    if not text:
        return
    last = 0
    for m in _INLINE_RE.finditer(text):
        start, end = m.span()
        if start > last:
            paragraph.add_run(text[last:start])
        token = m.group(1)
        if token.startswith("**") or token.startswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            # Single-char wrapping — italic.
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        last = end
    if last < len(text):
        paragraph.add_run(text[last:])


# ---- Block walker ---------------------------------------------------------


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[\-\*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_TABLE_ALIGN_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def _emit_table(doc, lines: list[str]) -> None:
    """Render a sequence of `| col | col |` lines as a docx table.
    Skips the alignment row (`|---|---|`). First content row becomes
    the header (bold). Cells are emitted with inline formatting."""
    rows: list[list[str]] = []
    for line in lines:
        if _TABLE_ALIGN_RE.match(line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell_text = row[c_idx] if c_idx < len(row) else ""
            # First paragraph of the cell already exists; reuse it.
            cell.paragraphs[0].text = ""
            _emit_inline_runs(cell.paragraphs[0], cell_text)
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def _markdown_to_docx_blocks(doc, md: str) -> None:
    """Walk the compiled markdown, dispatching each block to the right
    docx emitter. Stateful: tracks whether we're currently inside a
    list, a table, or a paragraph so consecutive lines don't fragment
    into one-line-per-paragraph output."""
    lines = md.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line — paragraph separator. Skip without emitting.
        if not stripped:
            i += 1
            continue

        # Heading
        m = _HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            style = (
                "Heading 1"
                if level == 1
                else "Heading 2"
                if level == 2
                else "Heading 3"
                if level == 3
                else "Heading 4"
            )
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                # Style not present in the empty default doc — fall
                # back to formatted-paragraph emission.
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
            _emit_inline_runs(p, text)
            i += 1
            continue

        # Table — collect contiguous `|...|` rows
        if _TABLE_ROW_RE.match(line):
            table_lines: list[str] = []
            while i < n and _TABLE_ROW_RE.match(lines[i]):
                table_lines.append(lines[i])
                i += 1
            _emit_table(doc, table_lines)
            continue

        # Block quote
        m = _BLOCKQUOTE_RE.match(stripped)
        if m:
            quote_lines: list[str] = []
            while i < n:
                bq = _BLOCKQUOTE_RE.match(lines[i].strip())
                if bq is None:
                    break
                quote_lines.append(bq.group(1))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            for run_text in [" ".join(quote_lines)]:
                run = p.add_run(run_text)
                run.italic = True
            continue

        # Bulleted list — collect consecutive bullets as separate
        # docx List Bullet paragraphs.
        m = _BULLET_RE.match(stripped)
        if m:
            while i < n:
                ln = lines[i].strip()
                bm = _BULLET_RE.match(ln)
                if bm is None:
                    break
                try:
                    p = doc.add_paragraph(style="List Bullet")
                except KeyError:
                    p = doc.add_paragraph()
                _emit_inline_runs(p, bm.group(1))
                i += 1
            continue

        # Numbered list
        m = _NUMBERED_RE.match(stripped)
        if m:
            while i < n:
                ln = lines[i].strip()
                nm = _NUMBERED_RE.match(ln)
                if nm is None:
                    break
                try:
                    p = doc.add_paragraph(style="List Number")
                except KeyError:
                    p = doc.add_paragraph()
                _emit_inline_runs(p, nm.group(1))
                i += 1
            continue

        # Plain paragraph — collect contiguous non-blank, non-special
        # lines as a single paragraph (markdown soft-wraps within a
        # paragraph the way HTML does).
        para_lines: list[str] = []
        while i < n:
            ln = lines[i]
            if not ln.strip():
                break
            if (
                _HEADING_RE.match(ln.strip())
                or _BULLET_RE.match(ln.strip())
                or _NUMBERED_RE.match(ln.strip())
                or _BLOCKQUOTE_RE.match(ln.strip())
                or _TABLE_ROW_RE.match(ln)
            ):
                break
            para_lines.append(ln.strip())
            i += 1
        p = doc.add_paragraph()
        _emit_inline_runs(p, " ".join(para_lines))


# ---- Document setup -------------------------------------------------------


def _apply_default_styles(doc: Document) -> None:
    """Tune the default Normal + heading styles to a reviewer-friendly
    federal-proposal look (Calibri 11pt body, Calibri 14/13/12 H1-H3,
    1" margins). Anything Word's defaults already do well stays."""
    # Margins.
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Body font.
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
    except KeyError:
        pass

    # Headings — keep Word's bold/spacing defaults; just nudge size +
    # color so they read as proposal headers, not slide titles.
    for style_name, size_pt, color in (
        ("Heading 1", 16, RGBColor(0x1F, 0x3A, 0x5F)),  # navy
        ("Heading 2", 13, RGBColor(0x1F, 0x3A, 0x5F)),
        ("Heading 3", 12, RGBColor(0x33, 0x33, 0x33)),
        ("Heading 4", 11, RGBColor(0x33, 0x33, 0x33)),
    ):
        try:
            s = doc.styles[style_name]
            s.font.name = "Calibri"
            s.font.size = Pt(size_pt)
            s.font.color.rgb = color
        except KeyError:
            continue


def _slug(text: str) -> str:
    """Filename-safe slug from arbitrary text."""
    import re as _re

    s = (text or "").strip().lower()
    s = _re.sub(r"[^a-z0-9]+", "-", s)
    return s.strip("-") or "proposal"


# ---- RFP-mandated filename extraction ------------------------------------

# Pre-filter: keywords that suggest a filename instruction is present
# in a submission_format compliance item. Cheap regex check so we don't
# pay for a Haiku call on every download when no convention exists
# (the common case for many RFPs).
_FILENAME_KEYWORDS_RE = re.compile(
    r"\b(?:file\s*name|filename|"
    r"(?:file|document)\s+(?:shall|must|should)\s+be\s+(?:named|saved|titled)|"
    r"name\s+(?:your|the)\s+(?:file|document|submission)|"
    r"save\s+(?:as|the\s+file)|"
    r"label\s+the\s+file|"
    r"naming\s+(?:convention|requirement)|"
    r"\.docx|\.pdf)\b",
    re.IGNORECASE,
)


_FILENAME_SYSTEM = (
    "You extract the REQUIRED output filename from RFP submission "
    "instructions. If the RFP specifies how the proposal file should "
    "be named, return ONLY the literal filename with placeholders "
    "filled in. If the RFP does not specify a filename, return null. "
    "Output ONLY a JSON object with one key: 'filename'."
)


def _build_filename_prompt(
    items_text: str,
    vendor: str,
    solicitation: str | None,
    agency: str | None,
    title: str | None,
) -> str:
    return f"""From the RFP submission instructions below, extract the required filename for the technical proposal DOCX (the main response document).

RULES:
- Substitute Vendor / RFP-Number / Agency placeholders with the actual values provided.
- Strip any directory or path separators — return just the bare filename.
- Append ".docx" if the convention specifies a name without extension.
- If the RFP does NOT specify a filename, return {{"filename": null}}.
- If multiple filenames are mentioned (e.g., separate technical + cost files), return the one for the main / technical proposal response.

VENDOR (legal name): {vendor}
SOLICITATION NUMBER: {solicitation or "(not stated — leave the placeholder if convention requires it)"}
AGENCY: {agency or "(not stated)"}
PROPOSAL TITLE: {title or "(not stated)"}

RFP SUBMISSION INSTRUCTIONS (verbatim quotes — pick the convention if any):
{items_text}

Return ONLY the JSON object."""


def _sanitize_filename(name: str) -> str:
    """Strip path separators, control chars, and Windows-forbidden
    characters; ensure a .docx extension; cap length. Defensive — the
    LLM should already produce a clean name, but we don't trust it
    blindly when the result becomes a filesystem path."""
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name)
    name = name.strip(". ")
    if not name:
        return ""
    if not name.lower().endswith(".docx"):
        name = f"{name}.docx"
    return name[:200]


def extract_submission_filename(proposal_id: int) -> str | None:
    """Inspect submission_format compliance items for an explicit
    output filename convention. When found, run a cheap Haiku call to
    fill in vendor / solicitation / agency placeholders and return the
    final filename. Returns None when no convention is stated — the
    common case — without spending an LLM call.

    Best-effort. Any failure (LLM error, parse failure, no profile)
    returns None so the caller falls back to its default filename.
    """
    import json as _json

    from app.core.company_profile import get_company_profile
    from app.db.session import SessionLocal
    from app.models import ComplianceMatrixItem, Proposal

    with SessionLocal() as db:
        proposal = db.get(Proposal, proposal_id)
        if proposal is None:
            return None
        items = (
            db.query(ComplianceMatrixItem)
            .filter(
                ComplianceMatrixItem.proposal_id == proposal_id,
                ComplianceMatrixItem.requirement_type == "submission_format",
                ComplianceMatrixItem.status == "active",
            )
            .all()
        )
        proposal_title = proposal.title
        agency = proposal.agency
        notes = proposal.notes or ""
        item_texts = [(it.requirement_text or "") for it in items]

    matches = [text for text in item_texts if _FILENAME_KEYWORDS_RE.search(text)]
    if not matches:
        return None

    items_text = "\n\n".join(f"- {t}" for t in matches[:20])

    try:
        profile = get_company_profile()
    except Exception:
        log.exception(
            "submission_filename: failed to load company profile",
        )
        return None
    vendor = (profile.get("company") or {}).get("legal_name") or "Vendor"

    # Solicitation number is parked in the proposal.notes field as
    # "Solicitation #: ..." by the intake-metadata auto-fill flow.
    solicitation: str | None = None
    sol_match = re.search(
        r"Solicitation\s*#:\s*([^\r\n]+)",
        notes,
        flags=re.IGNORECASE,
    )
    if sol_match:
        solicitation = sol_match.group(1).strip() or None

    from app.config import get_settings
    from app.services.llm import fmt_llm_usage, get_anthropic

    settings = get_settings()
    client = get_anthropic()
    prompt = _build_filename_prompt(
        items_text,
        vendor,
        solicitation,
        agency,
        proposal_title,
    )

    try:
        response, usage = client.complete(
            model=settings.model_light_extraction,
            system=_FILENAME_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=200,
            agent_name="submission_filename",
            proposal_id=proposal_id,
            temperature=0.0,
        )
    except Exception:
        log.exception(
            "submission_filename: Haiku call failed for proposal %d",
            proposal_id,
        )
        return None
    log.info(
        "submission_filename: proposal %d %s",
        proposal_id,
        fmt_llm_usage(usage),
    )

    obj_match = re.search(r"\{.*\}", response, re.DOTALL)
    if not obj_match:
        return None
    try:
        data = _json.loads(obj_match.group(0))
    except _json.JSONDecodeError:
        log.warning(
            "submission_filename: JSON parse failed for proposal %d: %r",
            proposal_id,
            response[:200],
        )
        return None
    raw = data.get("filename")
    if not raw or not isinstance(raw, str):
        return None
    cleaned = _sanitize_filename(raw.strip())
    return cleaned or None


# ---- Submission Checklist appendix ---------------------------------------

# Unicode ballot-box characters render reliably in Word and Google Docs.
# Bigger and more legible than plain "[ ]" / "[X]" in the rendered DOCX.
_BOX_OPEN = "☐"
_BOX_CHECKED = "☒"


def _emit_submission_checklist_appendix(
    doc: Document,
    snapshot: dict,
) -> None:
    """Render the Submission Checklist as an appendix in the DOCX.
    Page break before the appendix so it never bleeds into the last
    proposal section. Three subsections (RFP-required, User-tracked,
    System-verified) — each only emitted when it has rows.
    """
    totals = snapshot.get("totals", {})

    # Page break + appendix header. add_page_break() inserts the
    # break inside the previous paragraph so the next paragraph
    # starts on a fresh page; that's what we want for an appendix.
    last_para = doc.paragraphs[-1] if doc.paragraphs else None
    if last_para is not None:
        last_para.add_run().add_break()
    try:
        h1 = doc.add_paragraph(style="Heading 1")
    except KeyError:
        h1 = doc.add_paragraph()
    h1.add_run("Appendix: Submission Checklist")

    # Intro paragraph.
    intro = doc.add_paragraph()
    intro.add_run(
        "This appendix lists every item the submitter must obtain or "
        "address before submitting this proposal package. It is NOT "
        "part of the proposal narrative — included as a working "
        "checklist for the submission team."
    )

    # Banner with totals.
    pending = totals.get("all_obtained_pending", 0)
    banner = doc.add_paragraph()
    banner_run = banner.add_run(
        f"Status: {pending} item"
        f"{'s' if pending != 1 else ''} pending across all categories. "
        f"RFP-required forms/certs: "
        f"{totals.get('rfp_required_obtained', 0)} of "
        f"{totals.get('rfp_required_total', 0)} obtained · "
        f"User-tracked commitments: "
        f"{totals.get('commitments_obtained', 0)} of "
        f"{totals.get('commitments_total', 0)} obtained · "
        f"System-verified readiness: "
        f"{totals.get('system_checks_verified', 0)} of "
        f"{totals.get('system_checks_total', 0)} verified."
    )
    banner_run.italic = True
    banner.paragraph_format.space_after = Pt(8)

    # ---- RFP-required forms & certifications -----------------------
    rfp_required = snapshot.get("rfp_required") or []
    if rfp_required:
        try:
            h2 = doc.add_paragraph(style="Heading 2")
        except KeyError:
            h2 = doc.add_paragraph()
        h2.add_run(f"RFP-Required Forms & Certifications ({len(rfp_required)})")
        sub = doc.add_paragraph()
        sub_run = sub.add_run(
            "Auto-extracted from the RFP's compliance matrix. Each "
            "row needs an artifact attached or a written acknowledgement "
            "before submission."
        )
        sub_run.italic = True

        # Table: status / REQ ID / type / description / source / notes
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for col_idx, header_text in enumerate(
            [
                "Status",
                "REQ ID",
                "Type / Category",
                "Description",
                "Source",
            ]
        ):
            cell = hdr[col_idx]
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(header_text)
            run.bold = True

        for r in rfp_required:
            row = table.add_row().cells
            box = _BOX_CHECKED if r["obtained"] else _BOX_OPEN
            status_text = box + (" obtained" if r["obtained"] else "")
            row[0].text = status_text
            row[1].text = r["requirement_id"] or ""
            row[2].text = f"{r['requirement_type']} / {r['category']}"
            # Truncate to keep the table from running off the page.
            desc = (r.get("description") or "").strip()
            if len(desc) > 320:
                desc = desc[:317] + "…"
            row[3].text = desc
            src_bits: list[str] = []
            if r.get("source_section"):
                src_bits.append(f"§{r['source_section']}")
            if r.get("source_page"):
                src_bits.append(f"p.{r['source_page']}")
            row[4].text = " ".join(src_bits) or "—"
            if r.get("notes"):
                # Append notes as a second paragraph in the description
                # cell so they're visible in-context.
                p = row[3].add_paragraph()
                run = p.add_run(f"Notes: {r['notes']}")
                run.italic = True

    # ---- User-tracked commitments ---------------------------------
    commitments = snapshot.get("user_commitments") or []
    if commitments:
        try:
            h2 = doc.add_paragraph(style="Heading 2")
        except KeyError:
            h2 = doc.add_paragraph()
        h2.add_run(f"User-Tracked Commitments ({len(commitments)})")
        sub = doc.add_paragraph()
        sub_run = sub.add_run(
            "Items the user added during proposal development "
            "(post-resolution commitments, deliverables to attach, "
            "etc.)."
        )
        sub_run.italic = True

        for c in commitments:
            p = doc.add_paragraph()
            box = _BOX_CHECKED if c["obtained"] else _BOX_OPEN
            run = p.add_run(f"{box}  {c['description']}")
            if c["obtained"]:
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            if c.get("notes"):
                np = doc.add_paragraph()
                np_run = np.add_run(f"  Notes: {c['notes']}")
                np_run.italic = True
                np.paragraph_format.left_indent = Inches(0.3)

    # ---- System-verified readiness --------------------------------
    system_checks = snapshot.get("system_checks") or []
    if system_checks:
        try:
            h2 = doc.add_paragraph(style="Heading 2")
        except KeyError:
            h2 = doc.add_paragraph()
        h2.add_run(
            f"System-Verified Readiness Checks "
            f"({totals.get('system_checks_verified', 0)} of "
            f"{totals.get('system_checks_total', 0)} verified)"
        )
        sub = doc.add_paragraph()
        sub_run = sub.add_run(
            "Automatically computed from the proposal state. Use to "
            "confirm internal readiness before clicking Submit."
        )
        sub_run.italic = True

        for s in system_checks:
            p = doc.add_paragraph()
            box = _BOX_CHECKED if s["verified"] else _BOX_OPEN
            run = p.add_run(f"{box}  {s['label']} — {s.get('detail', '')}")
            if s["verified"]:
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            elif s.get("severity") == "critical":
                run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)


def compile_proposal_to_docx(
    proposal_id: int,
    *,
    include_cost_deferred: bool = True,
    include_submission_checklist: bool = True,
    proposal_title: str | None = None,
) -> tuple[bytes, str, dict[str, Any]]:
    """Compile the proposal's section drafts into a single DOCX byte
    stream. Returns (bytes, suggested_filename, summary_dict).

    `include_submission_checklist` (default True) appends an "Appendix:
    Submission Checklist" section after the main proposal content with
    every RFP-required form/certification, every user-tracked
    commitment, and every system-verified readiness check. The
    submitter uses this as the working checklist between download and
    actual submission. Pass False to suppress (e.g., for an export
    intended only for an evaluator's eyes).

    summary_dict mirrors the compiler's metadata so callers can
    surface skipped-sections / total-chars / checklist-pending
    counts in the UI without re-querying.

    Caller is responsible for serving the bytes (NiceGUI:
    `ui.download.content(bytes, filename)`; FastAPI:
    `Response(content=bytes, headers=...)`).
    """
    payload = compile_proposal_markdown(
        proposal_id,
        include_cost_deferred=include_cost_deferred,
    )
    md = payload["markdown"]

    doc = Document()
    _apply_default_styles(doc)

    if proposal_title:
        title_p = doc.add_paragraph()
        run = title_p.add_run(proposal_title.strip())
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        # Empty paragraph for spacing under the title.
        doc.add_paragraph()

    _markdown_to_docx_blocks(doc, md)

    checklist_summary: dict | None = None
    if include_submission_checklist:
        try:
            snapshot = get_submission_checklist_snapshot(proposal_id)
            _emit_submission_checklist_appendix(doc, snapshot)
            checklist_summary = snapshot.get("totals") or {}
        except Exception:
            log.exception(
                "compile_proposal_to_docx: failed to render submission "
                "checklist appendix for proposal %d — exporting without "
                "it (non-fatal).",
                proposal_id,
            )

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    filename = f"{_slug(proposal_title or f'proposal-{proposal_id}')}.docx"
    # Honor any RFP-mandated naming convention (e.g., the buyer
    # specifies "Files shall be named: VENDOR_RFP-NNN_TechProposal.docx").
    # Best-effort — falls through to the slug-based default on any
    # extraction failure or when no convention is stated.
    rfp_filename: str | None = None
    try:
        rfp_filename = extract_submission_filename(proposal_id)
    except Exception:
        log.exception(
            "compile_proposal_to_docx: filename extraction crashed for "
            "proposal %d — using default slug filename.",
            proposal_id,
        )
    if rfp_filename:
        log.info(
            "compile_proposal_to_docx: using RFP-mandated filename %r (default would have been %r)",
            rfp_filename,
            filename,
        )
        filename = rfp_filename
    summary = {
        "filename": filename,
        "filename_source": "rfp" if rfp_filename else "default",
        "byte_count": len(data),
        "total_sections": payload["total_sections"],
        "sections_included": payload["sections_included"],
        "sections_skipped": payload["sections_skipped"],
        "total_chars": payload["total_chars"],
        "submission_checklist": checklist_summary,
    }
    log.info(
        "compile_proposal_to_docx: proposal %d -> %d bytes, %d "
        "section(s) included, %d skipped, submission_checklist=%s",
        proposal_id,
        len(data),
        payload["total_sections"],
        len(payload["sections_skipped"]),
        checklist_summary or "skipped",
    )
    return data, filename, summary


__all__ = [
    "compile_proposal_to_docx",
]
